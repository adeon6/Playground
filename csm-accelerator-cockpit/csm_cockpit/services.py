from __future__ import annotations

import json
import hashlib
import os
import re
import shutil
import subprocess
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document


APP_VERSION = "0.6.0"

PROJECT_ROOT = Path(__file__).resolve().parents[1]
COCKPIT_ROOT = Path(__file__).resolve().parent
DEFAULT_CANONICAL_RUNS_DIR = Path(os.environ.get("USERPROFILE", str(Path.home()))) / "Documents" / "CODEX" / "alteryx" / "accelerator_projects"
RUNS_DIR = Path(os.environ.get("CODEX_ACCELERATOR_PROJECT_ROOT", str(DEFAULT_CANONICAL_RUNS_DIR))).expanduser()
LEGACY_RUNS_DIR = COCKPIT_ROOT / "runs"
STARTER_DOCS_DIR = PROJECT_ROOT / "starter_docs"
PROCESS_PACK_DIR = PROJECT_ROOT / "process_pack"
TOOLING_DIR = PROJECT_ROOT / "tooling"
SEQUENCE_CONFIG_PATH = PROCESS_PACK_DIR / "sequencer" / "sequence_config.json"
LARGE_PROMPT_PATH = PROCESS_PACK_DIR / "03_workflow_build" / "sop_to_alteryx_super_prompt_pack.md"
QUESTION_TEMPLATE_DOCX = PROCESS_PACK_DIR / "01_discovery" / "guided_discovery_conversation_template.docx"
DEFAULT_TRANSCRIPT_PATH = PROJECT_ROOT / "sample_data" / "sanitized_geo_spatial_discovery.md"
WORKFLOW_PROMPT_ARTIFACT = "status/codex_workflow_build_prompt.md"
WORKFLOW_HELPER_ARTIFACT = "status/START_CODEX_WORKFLOW_BUILD.ps1"
WORKFLOW_MANIFEST_ARTIFACT = "status/workflow_build_manifest.json"
LEGACY_PROMPT_ARTIFACT = "status/next_stage_prompt.md"
PIPELINE_STATUS_ARTIFACT = "status/pipeline_status.json"
PIPELINE_LOG_ARTIFACT = "status/pipeline_log.md"
PEER_REVIEW_STATUS_ARTIFACT = "status/peer_review_status.json"

CAPTURE_STATUSES = {
    "not_set": "Not selected",
    "answered": "Answered",
    "partial": "Partially answered",
    "not_answered": "Not answered",
    "needs_follow_up": "Needs follow-up",
}

EVIDENCE_STATUSES = {
    "supported": "Answered by transcript",
    "weak_evidence": "Partially answered by transcript",
    "missing": "Not answered by transcript",
    "conflicting": "Needs review",
    "not_run": "Not answered by transcript",
}

DOC_ARTIFACTS = [
    "01_customer_discovery_conversation.md",
    "02_guided_sop_capture.md",
    "03_accelerator_sop.md",
    "sop_gap_log.md",
    "sop_architecture_assessment.md",
]

ACCELERATOR_ASSET_ARTIFACTS = [
    "04_value_statement.md",
    "05_use_case_summary.md",
    "06_case_study_skeleton.md",
    "07_accelerator_101.md",
    "08_accelerator_102.md",
    "09_accelerator_201.md",
]

ARTIFACT_LABELS = {
    "01_customer_discovery_conversation.md": "Discovery conversation",
    "02_guided_sop_capture.md": "Guided SOP capture",
    "03_accelerator_sop.md": "Accelerator SOP",
    "sop_gap_log.md": "Gap log",
    "sop_architecture_assessment.md": "Architecture assessment",
    "04_value_statement.md": "Value statement",
    "05_use_case_summary.md": "Use-case summary",
    "06_case_study_skeleton.md": "Case-study skeleton",
    "07_accelerator_101.md": "Accelerator 101",
    "08_accelerator_102.md": "Accelerator 102",
    "09_accelerator_201.md": "Accelerator 201",
    WORKFLOW_PROMPT_ARTIFACT: "Codex workflow build prompt",
    WORKFLOW_HELPER_ARTIFACT: "Codex launch helper",
    WORKFLOW_MANIFEST_ARTIFACT: "Workflow build manifest",
    LEGACY_PROMPT_ARTIFACT: "Workflow handoff prompt",
    "status/pipeline_status.json": "Pipeline status",
    "status/pipeline_log.md": "Pipeline log",
    PEER_REVIEW_STATUS_ARTIFACT: "Peer-review status",
}

JON_PROCESS_FOLDERS = [
    "00_start_here",
    "01_discovery",
    "02_sop_authoring",
    "03_workflow_build",
    "04_reference_examples",
    "sequencer",
]

WORKFLOW_STATUS_DIR = "03_workflow_build/status"
WORKFLOW_OUTPUT_DIR = "03_workflow_build/workflows"
WORKFLOW_VALIDATION_DIR = "03_workflow_build/validation"
ACCELERATOR_ASSET_DIR = "04_reference_examples/accelerator_assets"

ARTIFACT_PATHS = {
    "01_customer_discovery_conversation.md": "01_discovery/01_customer_discovery_conversation.md",
    "02_guided_sop_capture.md": "02_sop_authoring/02_guided_sop_capture.md",
    "03_accelerator_sop.md": "02_sop_authoring/03_accelerator_sop.md",
    "sop_gap_log.md": "02_sop_authoring/sop_gap_log.md",
    "sop_architecture_assessment.md": "02_sop_authoring/sop_architecture_assessment.md",
    "04_value_statement.md": f"{ACCELERATOR_ASSET_DIR}/04_value_statement.md",
    "05_use_case_summary.md": f"{ACCELERATOR_ASSET_DIR}/05_use_case_summary.md",
    "06_case_study_skeleton.md": f"{ACCELERATOR_ASSET_DIR}/06_case_study_skeleton.md",
    "07_accelerator_101.md": f"{ACCELERATOR_ASSET_DIR}/07_accelerator_101.md",
    "08_accelerator_102.md": f"{ACCELERATOR_ASSET_DIR}/08_accelerator_102.md",
    "09_accelerator_201.md": f"{ACCELERATOR_ASSET_DIR}/09_accelerator_201.md",
    PEER_REVIEW_STATUS_ARTIFACT: f"{ACCELERATOR_ASSET_DIR}/peer_review_status.json",
    WORKFLOW_PROMPT_ARTIFACT: f"{WORKFLOW_STATUS_DIR}/codex_workflow_build_prompt.md",
    WORKFLOW_HELPER_ARTIFACT: f"{WORKFLOW_STATUS_DIR}/START_CODEX_WORKFLOW_BUILD.ps1",
    WORKFLOW_MANIFEST_ARTIFACT: f"{WORKFLOW_STATUS_DIR}/workflow_build_manifest.json",
    LEGACY_PROMPT_ARTIFACT: f"{WORKFLOW_STATUS_DIR}/next_stage_prompt.md",
    PIPELINE_STATUS_ARTIFACT: f"{WORKFLOW_STATUS_DIR}/pipeline_status.json",
    PIPELINE_LOG_ARTIFACT: f"{WORKFLOW_STATUS_DIR}/pipeline_log.md",
}

PROJECT_FOLDERS = [
    *JON_PROCESS_FOLDERS,
    "03_workflow_build/status",
    "03_workflow_build/workflows",
    "03_workflow_build/validation",
    "04_reference_examples/accelerator_assets",
    "data/raw/transcripts",
    "data/generated",
    "tooling",
]

VISIBLE_ARTIFACTS = [
    *DOC_ARTIFACTS,
    *ACCELERATOR_ASSET_ARTIFACTS,
    PEER_REVIEW_STATUS_ARTIFACT,
    WORKFLOW_PROMPT_ARTIFACT,
    WORKFLOW_HELPER_ARTIFACT,
    WORKFLOW_MANIFEST_ARTIFACT,
]

ALL_ARTIFACTS = [
    *VISIBLE_ARTIFACTS,
    LEGACY_PROMPT_ARTIFACT,
    PIPELINE_STATUS_ARTIFACT,
    PIPELINE_LOG_ARTIFACT,
]

REQUIRED_FOR_SOP_GATE = {
    "business_problem",
    "current_process",
    "desired_outcome",
    "value_realization",
    "business_questions",
    "scope",
    "inputs_sources_ownership",
    "rules_logic_definitions",
    "exceptions_safe_handling",
    "validation_trust",
    "operational_readiness_phasing",
}

UNCERTAINTY_TERMS = [
    "not sure",
    "don't know",
    "do not know",
    "unknown",
    "tbc",
    "to be confirmed",
    "haven't decided",
    "have not decided",
    "need to check",
    "need follow",
    "we need to ask",
]


@dataclass(frozen=True)
class Section:
    id: str
    label: str
    prompt: str
    questions: list[str]
    capture_fields: list[str]
    keywords: list[str]
    required: bool = True
    gate_reason: str = "Required for SOP and workflow-build handoff."


JON_SECTIONS = [
    Section(
        "business_problem",
        "Business Problem",
        "Capture the pain, business consequence, affected users, and why solving it matters now.",
        [
            "What problem are you trying to solve?",
            "What is slow, manual, inconsistent, risky, or hard to trust today?",
            "What happens when this problem is not solved well?",
            "Who feels the pain most directly?",
            "Are you trying to improve visibility, speed, quality, prioritisation, decision-making, or something else?",
        ],
        ["Core problem", "Business impact", "Affected teams", "Why now", "Cost of current state"],
        ["problem", "pain", "challenge", "manual", "slow", "risk", "impact", "cost", "inconsistent", "trust"],
    ),
    Section(
        "current_process",
        "Current Process",
        "Describe the current-state process, handoffs, systems, spreadsheets, and known friction.",
        [
            "How is this handled today?",
            "What are the main steps from input to outcome?",
            "Which systems, files, reports, or teams are involved?",
            "Where do manual work, spreadsheets, emails, or handoffs happen?",
            "Which steps are most painful or fragile?",
            "What makes the current output difficult to trust or use?",
        ],
        ["Current process", "Systems involved", "Manual steps", "Handoffs", "Bottlenecks"],
        ["current process", "today", "as is", "manual steps", "spreadsheet", "handoff", "workflow", "process"],
    ),
    Section(
        "desired_outcome",
        "Desired Outcome",
        "Clarify what first useful output should exist and what business decision or action it enables.",
        [
            "If we gave you a useful first version quickly, what would it produce?",
            "What outputs would be most valuable first?",
            "Who would use those outputs?",
            "What action or decision should those outputs support?",
            "What does success look like for the first slice?",
            "What can wait until later?",
        ],
        ["Target output", "Consumer", "Decision / action", "First-slice success", "Deferred items"],
        ["output", "deliverable", "dashboard", "report", "app", "action", "decision", "success", "first version"],
    ),
    Section(
        "value_realization",
        "Value Realisation",
        "Capture how the customer will recognise value, measure improvement, and justify the accelerator.",
        [
            "What business value should this first slice create if it works well?",
            "Where is the value likely to show up first: time saved, faster decisions, reduced risk, improved consistency, better visibility, or greater trust?",
            "Are there any quantitative signals we can capture now, such as effort, turnaround time, backlog, volume, or error rate?",
            "If the value is not yet measurable, what qualitative improvement would still matter?",
            "What would make stakeholders say this is worth continuing or expanding?",
            "What value assumptions still need to be validated later?",
        ],
        [
            "Current pain or business impact",
            "Workflow-enabled operational change",
            "Quantitative value signals",
            "Qualitative value signals",
            "Trust / adoption signals",
            "Value assumptions / unknowns to validate later",
        ],
        [
            "value",
            "value realization",
            "value realisation",
            "benefit",
            "roi",
            "saving",
            "baseline",
            "measure",
            "metric",
            "improvement",
            "quantify",
            "quantify success",
            "success",
            "time saved",
            "time taken",
            "halve",
            "reduce",
            "reduction",
            "faster",
            "quickly",
            "efficiency",
        ],
    ),
    Section(
        "business_questions",
        "Business Questions",
        "List the questions the accelerator must answer and how users decide what to do next.",
        [
            "What questions do you need this solution to answer?",
            "What should someone be able to decide once they see the results?",
            "Which cases should be prioritised?",
            "Which cases should be flagged for review?",
            "What patterns, exceptions, or risks matter most?",
        ],
        ["Core questions", "Decision points", "Prioritisation", "Review cases", "Business significance"],
        ["question", "decide", "prioritize", "prioritise", "rank", "exception", "review", "what should", "which cases"],
    ),
    Section(
        "scope",
        "Scope",
        "Separate the first slice from explicit exclusions, pilot boundaries, and later phases.",
        [
            "What is in scope for the first slice?",
            "What is out of scope for now?",
            "Are there specific business units, regions, products, customers, channels, or process steps to include or exclude?",
            "Is there a preferred pilot boundary?",
            "Are we proving feasibility, value, or both?",
        ],
        ["In-scope entities / domains", "Out-of-scope areas", "Pilot boundary", "Phase 1 intent", "Explicit exclusions"],
        ["scope", "in scope", "out of scope", "pilot", "phase", "priority", "region", "business unit", "product"],
    ),
    Section(
        "inputs_sources_ownership",
        "Inputs, Sources, And Ownership",
        "Identify source systems, extract owners, refresh cadence, access constraints, and package-safe inputs.",
        [
            "What information is needed to answer these questions?",
            "Where does that information live today?",
            "Which systems, files, owners, or teams provide it?",
            "What do we know about the source structure already?",
            "What still needs to be confirmed before build?",
            "Are there known quality, access, or timing concerns?",
        ],
        [
            "Required inputs",
            "Source systems / files",
            "Data owners / SMEs",
            "Known structure and availability",
            "Unknowns requiring discovery",
            "Quality / access risks",
        ],
        ["source", "system", "input", "file", "table", "extract", "owner", "refresh", "cadence", "access", "field", "column", "entity", "grain", "join", "key", "quality"],
    ),
    Section(
        "rules_logic_definitions",
        "Rules, Logic, And Definitions",
        "Document definitions, calculations, thresholds, filters, joins, and exception-handling rules.",
        [
            "What business rules are already trusted?",
            "Which metrics, classifications, or thresholds matter?",
            "How do you currently define a good, bad, at-risk, or priority case?",
            "Are there rules for inclusion, exclusion, or segmentation?",
            "Which edge cases need special treatment?",
            "What assumptions are acceptable, and which require confirmation?",
        ],
        ["Trusted rules", "Important definitions", "Thresholds / classifications", "Inclusion / exclusion rules", "Edge cases", "Areas where logic is still uncertain"],
        ["rule", "logic", "definition", "threshold", "join", "filter", "criteria", "calculation", "formula"],
    ),
    Section(
        "exceptions_safe_handling",
        "Exceptions And Safe Handling",
        "Capture how missing, inconsistent, late, invalid, or risky data should be handled safely.",
        [
            "What should happen when data is missing, inconsistent, late, or invalid?",
            "What failure cases do you already expect?",
            "Should questionable records be excluded, flagged, defaulted, or sent for review?",
            "What outcomes would create risk if handled incorrectly?",
        ],
        ["Expected failure modes", "Safe handling rules", "Review / escalation paths", "Risk-sensitive cases"],
        ["exception", "missing", "inconsistent", "late", "invalid", "failure", "exclude", "flag", "default", "review", "risk", "safe"],
    ),
    Section(
        "validation_trust",
        "Validation And Trust",
        "Define how outputs are reconciled, who validates them, and what acceptance checks prove trust.",
        [
            "What would make you trust the first slice?",
            "How would you validate whether it is working?",
            "What should it reconcile against?",
            "What should be explainable to the business?",
            "What would count as failure?",
            "Do different audiences need different outputs from the same logic?",
        ],
        ["Validation criteria", "Reconciliation targets", "Trust signals", "Explainability needs", "Failure conditions", "Audience-specific output needs"],
        ["validate", "validation", "reconcile", "check", "test", "trust", "review", "sign-off", "compare", "explainable", "failure"],
    ),
    Section(
        "operational_readiness_phasing",
        "Operational Readiness And Phasing",
        "Capture what is needed for short-term usefulness versus later productionisation.",
        [
            "What needs to be true for this to be useful in the short term?",
            "What can be manual in phase 1?",
            "What needs to be productionised later?",
            "Are there scheduling, alerting, write-back, governance, or access requirements that should be deferred?",
            "What is the smallest slice that would still prove the approach?",
        ],
        ["Immediate needs", "Deferred operational requirements", "Phase 1 vs later-state expectations", "Minimum viable first slice"],
        ["operational", "readiness", "schedule", "refresh", "run", "owner", "handover", "deploy", "production", "support", "phase", "manual", "governance", "access"],
    ),
]


def slugify(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", value.strip()).strip("-").lower()
    return cleaned or "accelerator-project"


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def template_root() -> Path:
    return COCKPIT_ROOT / "templates"


def static_root() -> Path:
    return COCKPIT_ROOT / "static"


def load_sequence_config() -> dict[str, Any]:
    if SEQUENCE_CONFIG_PATH.exists():
        try:
            return json.loads(SEQUENCE_CONFIG_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass
    return {
        "project_structure": PROJECT_FOLDERS,
        "stages": [
            {"id": "01_discovery", "name": "Discovery"},
            {"id": "02_guided_capture", "name": "Guided SOP capture"},
            {"id": "03_accelerator_sop", "name": "Accelerator SOP"},
            {"id": "04_workflow_build", "name": "Workflow build handoff"},
            {"id": "05_review", "name": "Review"},
        ],
    }


def load_question_bank(template_path: Path = QUESTION_TEMPLATE_DOCX) -> list[Section]:
    # Jon's interview script is the source of truth for V3. The Word template is bundled
    # for reference, but the app keeps section-level prompts stable for CSM usability.
    return JON_SECTIONS


def read_docx_text(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_transcript_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(path)
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError("Transcript must be a .docx, .md, or .txt file.")


def split_evidence_units(text: str) -> list[str]:
    rough_units = re.split(r"[\r\n]{2,}|(?<=[.!?])\s+(?=[A-Z0-9])", text)
    units = []
    for unit in rough_units:
        cleaned = re.sub(r"\s+", " ", unit).strip()
        if len(cleaned) >= 35:
            units.append(cleaned)

    # Discovery transcripts often use short section-heading blocks:
    # "10. Value Realization" / "How would you quantify success?" / answer.
    # Those lines are individually too short for rough sentence scoring, so add
    # small rolling line windows to preserve the heading + Q/A context.
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if line.strip()]
    for start in range(len(lines)):
        window = " ".join(lines[start : start + 5]).strip()
        if len(window) >= 35:
            units.append(window)

    seen: set[str] = set()
    unique_units: list[str] = []
    for unit in units:
        key = unit.lower()
        if key not in seen:
            seen.add(key)
            unique_units.append(unit)
    return unique_units


def _keyword_hits(text: str, keywords: list[str]) -> tuple[int, list[str]]:
    total = 0
    found: list[str] = []
    lower_text = text.lower()
    for keyword in keywords:
        count = len(re.findall(re.escape(keyword.lower()), lower_text))
        if count:
            total += count
            found.append(keyword)
    return total, found


def _top_evidence(units: list[str], section: Section, limit: int = 3) -> list[dict[str, Any]]:
    scored: list[tuple[int, list[str], str]] = []
    for unit in units:
        score, keywords = _keyword_hits(unit, section.keywords)
        if score:
            scored.append((score, keywords, unit))
    scored.sort(key=lambda item: (item[0], len(item[2])), reverse=True)
    return [
        {
            "score": score,
            "keywords": keywords[:6],
            "snippet": snippet[:420],
        }
        for score, keywords, snippet in scored[:limit]
    ]


def _evidence_summary(section: Section, evidence: list[dict[str, Any]], status: str) -> str:
    if status == "missing":
        return "No useful transcript signal was found for this section."
    found_keywords = sorted({keyword for item in evidence for keyword in item.get("keywords", [])})
    keyword_text = ", ".join(found_keywords[:5]) if found_keywords else "section terms"
    if status == "supported":
        return f"Transcript evidence repeatedly touches this section through {keyword_text}."
    if status == "weak_evidence":
        return f"Transcript evidence hints at this section through {keyword_text}, but needs CSM confirmation."
    return "Transcript includes uncertainty or ambiguity that should be resolved before approval."


def analyze_transcript_text(
    text: str,
    sections: list[Section],
    capture: dict[str, dict[str, Any]] | None = None,
) -> dict[str, dict[str, Any]]:
    capture = capture or {}
    units = split_evidence_units(text)
    analysis: dict[str, dict[str, Any]] = {}
    for section in sections:
        evidence = _top_evidence(units, section)
        score = sum(item["score"] for item in evidence)
        status = "missing"
        if score >= 6 or len(evidence) >= 3:
            status = "supported"
        elif score >= 2:
            status = "weak_evidence"

        capture_status = capture.get(section.id, {}).get("status", "not_answered")
        combined_snippets = " ".join(item["snippet"].lower() for item in evidence)
        has_uncertainty = any(term in combined_snippets for term in UNCERTAINTY_TERMS)
        if has_uncertainty and capture_status in {"answered", "partial"} and status != "missing":
            status = "conflicting"

        recommendation = "Capture this as a follow-up, assumption, or confirmed gap before SOP approval."
        if status == "supported" and capture_status == "not_answered":
            recommendation = "Transcript appears to answer this; CSM should summarize and approve it."
        elif status == "supported":
            recommendation = "Evidence is strong enough for CSM playback and approval."
        elif status == "weak_evidence":
            recommendation = "Partial evidence found. CSM should confirm or add a follow-up."
        elif status == "conflicting":
            recommendation = "Transcript includes uncertainty that conflicts with the captured status."

        analysis[section.id] = {
            "status": status,
            "score": score,
            "summary": _evidence_summary(section, evidence, status),
            "evidence": evidence,
            "recommendation": recommendation,
            "updated_at": utc_now(),
        }

    return analysis


def ensure_runs_dir() -> None:
    RUNS_DIR.resolve().mkdir(parents=True, exist_ok=True)


def run_dir(run_id: str) -> Path:
    safe_run_id = slugify(run_id)
    if safe_run_id != run_id:
        raise ValueError("Invalid project id.")
    return (RUNS_DIR / run_id).resolve()


def manifest_path(run_id: str) -> Path:
    primary = run_dir(run_id) / "manifest.json"
    if primary.exists() or not LEGACY_RUNS_DIR.exists():
        return primary
    legacy = (LEGACY_RUNS_DIR / run_id / "manifest.json").resolve()
    if legacy.exists():
        return legacy
    return primary


def ensure_project_structure(project_dir: Path) -> None:
    for folder in PROJECT_FOLDERS:
        (project_dir / folder).mkdir(parents=True, exist_ok=True)
    _sync_process_pack_scaffold(project_dir)
    _sync_project_tooling(project_dir)


def _copytree_fresh(source: Path, target: Path) -> None:
    if not source.exists():
        return
    shutil.copytree(
        source,
        target,
        dirs_exist_ok=True,
        ignore=shutil.ignore_patterns(
            "__pycache__",
            ".pytest_cache",
            ".mypy_cache",
            ".venv",
            "venv",
            "node_modules",
            "*.pyc",
            "*.pyo",
            ".DS_Store",
        ),
    )


def _sync_project_tooling(project_dir: Path) -> None:
    """Keep each handoff project self-contained for local Codex."""
    tooling_target = project_dir / "tooling"
    _copytree_fresh(TOOLING_DIR / "alteryx_workflow_builder", tooling_target / "alteryx_workflow_builder")
    _copytree_fresh(TOOLING_DIR / "alteryx-beautification", tooling_target / "alteryx-beautification")


def _sync_process_pack_scaffold(project_dir: Path) -> None:
    """Seed Jon's operating-system folders without replacing generated run files."""
    for folder in JON_PROCESS_FOLDERS:
        _copytree_fresh(PROCESS_PACK_DIR / folder, project_dir / folder)
    if (PROCESS_PACK_DIR / "README.md").exists():
        start_here = project_dir / "00_start_here"
        start_here.mkdir(parents=True, exist_ok=True)
        shutil.copy2(PROCESS_PACK_DIR / "README.md", start_here / "README_process_pack.md")


def _artifact_path(project_dir: Path, artifact: str) -> Path:
    normalized = artifact.replace("\\", "/")
    mapped = ARTIFACT_PATHS.get(normalized)
    if mapped:
        return project_dir / mapped
    if normalized.startswith("status/"):
        return project_dir / WORKFLOW_STATUS_DIR / normalized.removeprefix("status/")
    return project_dir / normalized


def _artifact_record(project_dir: Path, artifact: str) -> dict[str, Any]:
    path = _artifact_path(project_dir, artifact)
    if artifact in DOC_ARTIFACTS:
        group = "internal_build_docs"
        group_label = "Internal build doc"
    elif artifact in ACCELERATOR_ASSET_ARTIFACTS:
        group = "customer_facing_assets"
        group_label = "Customer-facing asset"
    elif artifact.startswith("status/"):
        group = "status_and_review"
        group_label = "Status / review"
    else:
        group = "workflow_handoff"
        group_label = "Workflow handoff"
    return {
        "name": artifact,
        "label": ARTIFACT_LABELS.get(artifact, artifact),
        "path": str(path),
        "relative_path": str(path.relative_to(project_dir)).replace("\\", "/"),
        "exists": path.exists(),
        "status": "generated" if path.exists() else "not_generated",
        "group": group,
        "group_label": group_label,
    }


def _project_identity_hash(run_id: str, customer_name: str, project_name: str, canonical_project_root: str) -> str:
    payload = json.dumps(
        {
            "run_id": run_id,
            "customer_name": customer_name,
            "project_name": project_name,
            "canonical_project_root": str(Path(canonical_project_root).resolve()),
        },
        sort_keys=True,
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]


def _ensure_project_identity(manifest: dict[str, Any], project_dir: Path | None = None) -> dict[str, Any]:
    project_dir = (project_dir or run_dir(manifest["run_id"])).resolve()
    identity = manifest.setdefault("project_identity", {})
    identity["run_id"] = manifest["run_id"]
    identity["canonical_project_root"] = str(project_dir)
    identity["workspace_root"] = str(RUNS_DIR.resolve())
    identity["manifest_path"] = str(project_dir / "manifest.json")
    identity["identity_hash"] = _project_identity_hash(
        manifest["run_id"],
        manifest.get("customer_name", ""),
        manifest.get("project_name", ""),
        identity["canonical_project_root"],
    )
    manifest["canonical_project_root"] = identity["canonical_project_root"]
    return manifest


SECTION_KEY_MIGRATIONS = {
    "scope_priorities": "scope",
    "source_systems": "inputs_sources_ownership",
    "data_shape_entities": "inputs_sources_ownership",
    "business_rules": "rules_logic_definitions",
    "output_action": "desired_outcome",
    "operational_constraints": "operational_readiness_phasing",
}


def _merge_capture_item(existing: dict[str, Any], incoming: dict[str, Any], source_label: str) -> dict[str, Any]:
    merged = dict(existing)
    if not merged.get("status") or merged.get("status") in {"not_set", "not_answered"}:
        merged["status"] = incoming.get("status", merged.get("status", "not_answered"))
    existing_notes = str(merged.get("notes", "")).strip()
    incoming_notes = str(incoming.get("notes", "")).strip()
    if incoming_notes and incoming_notes not in existing_notes:
        merged["notes"] = f"{existing_notes}\n\nMigrated from {source_label}: {incoming_notes}".strip()
    else:
        merged["notes"] = existing_notes
    merged["approved"] = bool(merged.get("approved") or incoming.get("approved"))
    return merged


def migrate_manifest_sections(manifest: dict[str, Any], sections: list[Section] | None = None) -> dict[str, Any]:
    sections = sections or JON_SECTIONS
    valid_ids = {section.id for section in sections}
    capture = manifest.setdefault("capture", {})
    for section in sections:
        capture.setdefault(section.id, {"status": "not_set", "notes": "", "approved": False})
    for old_key, new_key in SECTION_KEY_MIGRATIONS.items():
        if old_key in capture:
            capture[new_key] = _merge_capture_item(capture.get(new_key, {}), capture[old_key], old_key)
            capture.pop(old_key, None)
    for key in list(capture.keys()):
        if key not in valid_ids:
            capture.pop(key, None)
            continue
        item = capture.setdefault(key, {})
        if item.get("status") not in CAPTURE_STATUSES:
            item["status"] = "not_set"
        item.setdefault("notes", "")
        item.setdefault("approved", False)

    analysis = manifest.setdefault("analysis", {})
    for old_key, new_key in SECTION_KEY_MIGRATIONS.items():
        if old_key in analysis and new_key not in analysis:
            analysis[new_key] = analysis[old_key]
        analysis.pop(old_key, None)
    for key in list(analysis.keys()):
        if key not in valid_ids:
            analysis.pop(key, None)
    return manifest


def new_manifest(customer_name: str, project_name: str, csm_name: str, sections: list[Section]) -> dict[str, Any]:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    base = slugify(customer_name or project_name or "customer")
    run_id = f"{timestamp}-{base}"
    project_dir = run_dir(run_id)
    clean_customer_name = customer_name.strip() or "Unnamed customer"
    clean_project_name = project_name.strip() or "Accelerator discovery"
    manifest = {
        "schema_version": 4,
        "run_id": run_id,
        "customer_name": clean_customer_name,
        "project_name": clean_project_name,
        "csm_name": csm_name.strip(),
        "created_at": utc_now(),
        "updated_at": utc_now(),
        "canonical_project_root": str(project_dir),
        "project_structure": PROJECT_FOLDERS,
        "process_pack": {
            "source": "Jon accelerator operating system",
            "large_prompt": str(LARGE_PROMPT_PATH),
            "sequence_config": str(SEQUENCE_CONFIG_PATH),
        },
        "transcript": {
            "source_path": "",
            "stored_path": "",
            "canonical_path": str(_artifact_path(project_dir, "01_customer_discovery_conversation.md")),
            "characters": 0,
            "analyzed_at": "",
        },
        "transcripts": [],
        "capture": {
            section.id: {
                "status": "not_set",
                "notes": "",
                "approved": False,
            }
            for section in sections
        },
        "analysis": {},
        "artifacts": {
            artifact: _artifact_record(project_dir, artifact)
            for artifact in ALL_ARTIFACTS
        },
        "workflow_build": {
            "status": "not_generated",
            "prompt_path": "",
            "helper_script_path": "",
            "manifest_path": "",
            "prompt_artifact": WORKFLOW_PROMPT_ARTIFACT,
            "helper_artifact": WORKFLOW_HELPER_ARTIFACT,
            "manifest_artifact": WORKFLOW_MANIFEST_ARTIFACT,
            "codex_detected": False,
            "codex_path": "",
            "codex_detection_note": "",
            "designer_detected": False,
            "engine_path": "",
            "designer_path": "",
            "tooling_bundle_version": "",
            "builder_toolkit_ready": False,
            "beautification_ready": False,
            "detected_artifacts": [],
            "generated_at": "",
            "last_launch_at": "",
            "last_launch_status": "",
        },
        "sync": {
            "synced_at": "",
            "target_docs_dir": str(STARTER_DOCS_DIR),
        },
    }
    return _ensure_project_identity(manifest, project_dir)


def _refresh_artifact_records(manifest: dict[str, Any]) -> dict[str, Any]:
    project_dir = run_dir(manifest["run_id"])
    _ensure_project_identity(manifest, project_dir)
    migrate_manifest_sections(manifest)
    records = manifest.setdefault("artifacts", {})
    for artifact in ALL_ARTIFACTS:
        existing = records.get(artifact, {})
        generated_at = existing.get("generated_at", "")
        records[artifact] = _artifact_record(project_dir, artifact)
        if generated_at and records[artifact]["exists"]:
            records[artifact]["generated_at"] = generated_at
    refresh_workflow_build_state(manifest)
    return manifest


def save_manifest(manifest: dict[str, Any]) -> None:
    ensure_runs_dir()
    project_dir = run_dir(manifest["run_id"])
    ensure_project_structure(project_dir)
    _ensure_project_identity(manifest, project_dir)
    migrate_manifest_sections(manifest)
    manifest["updated_at"] = utc_now()
    _refresh_artifact_records(manifest)
    path = manifest_path(manifest["run_id"])
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def load_manifest(run_id: str) -> dict[str, Any]:
    manifest = json.loads(manifest_path(run_id).read_text(encoding="utf-8"))
    _ensure_project_identity(manifest)
    migrate_manifest_sections(manifest)
    _transcript_records(manifest)
    return _refresh_artifact_records(manifest)


def list_manifests() -> list[dict[str, Any]]:
    ensure_runs_dir()
    manifests = []
    seen: set[str] = set()
    for root in [RUNS_DIR, LEGACY_RUNS_DIR]:
        if not root.exists():
            continue
        for path in sorted(root.glob("*/manifest.json"), reverse=True):
            try:
                manifest = json.loads(path.read_text(encoding="utf-8"))
                if manifest.get("run_id") in seen:
                    continue
                _ensure_project_identity(manifest)
                migrate_manifest_sections(manifest)
                _transcript_records(manifest)
                manifests.append(_refresh_artifact_records(manifest))
                seen.add(manifest.get("run_id", ""))
            except (json.JSONDecodeError, ValueError):
                continue
    return manifests


def update_capture_from_form(manifest: dict[str, Any], form: dict[str, Any], sections: list[Section]) -> dict[str, Any]:
    capture = manifest.setdefault("capture", {})
    for section in sections:
        item = capture.setdefault(section.id, {})
        status = str(form.get(f"status_{section.id}", item.get("status", "not_set")))
        item["status"] = status if status in CAPTURE_STATUSES else "not_set"
        item["notes"] = str(form.get(f"notes_{section.id}", item.get("notes", ""))).strip()
    return manifest


def reanalyze_transcript_corpus(manifest: dict[str, Any], sections: list[Section]) -> dict[str, Any]:
    if not _transcript_records(manifest):
        return manifest
    combined_text = _refresh_canonical_transcript(manifest)
    manifest["analysis"] = analyze_transcript_text(combined_text, sections, manifest.get("capture", {}))
    return manifest


def update_approvals_from_form(manifest: dict[str, Any], form: dict[str, Any], sections: list[Section]) -> dict[str, Any]:
    capture = manifest.setdefault("capture", {})
    for section in sections:
        item = capture.setdefault(section.id, {})
        item["approved"] = form.get(f"approved_{section.id}") == "on"
    return manifest


def _transcript_records(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    records = manifest.setdefault("transcripts", [])
    if records:
        return records
    legacy = manifest.get("transcript", {})
    stored_path = legacy.get("stored_path", "")
    if stored_path:
        records.append(
            {
                "id": "legacy-transcript",
                "source_type": "legacy",
                "source_path": legacy.get("source_path", ""),
                "original_name": Path(stored_path).name,
                "stored_path": stored_path,
                "text_path": "",
                "characters": legacy.get("characters", 0),
                "attached_at": legacy.get("analyzed_at", ""),
                "is_demo": False,
            }
        )
    return records


def _write_transcript_text_file(target: Path, text: str) -> Path:
    text_target = target.with_suffix(target.suffix + ".txt")
    text_target.write_text(text, encoding="utf-8")
    return text_target


def _unique_transcript_target(target_dir: Path, source_name: str) -> Path:
    source = Path(source_name)
    suffix = source.suffix.lower()
    stem = slugify(source.stem) or "transcript"
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = target_dir / f"{timestamp}-{stem}{suffix}"
    counter = 2
    while target.exists():
        target = target_dir / f"{timestamp}-{stem}-{counter}{suffix}"
        counter += 1
    return target


def _combined_transcript_text(manifest: dict[str, Any]) -> str:
    blocks: list[str] = []
    for index, record in enumerate(_transcript_records(manifest), start=1):
        stored_path = Path(record.get("stored_path", ""))
        text_path = Path(record.get("text_path", "")) if record.get("text_path") else None
        text = ""
        if text_path and text_path.exists():
            text = text_path.read_text(encoding="utf-8", errors="replace")
        elif stored_path.exists():
            text = read_transcript_text(stored_path)
        if not text.strip():
            continue
        original_name = record.get("original_name") or stored_path.name or f"transcript-{index}"
        blocks.append(
            "\n".join(
                [
                    f"## Transcript Source {index}: {original_name}",
                    "",
                    f"- Source type: {record.get('source_type', 'unknown')}",
                    f"- Attached: {record.get('attached_at', 'unknown')}",
                    f"- Transcript ID: `{record.get('id', f'transcript-{index}')}`",
                    "",
                    text.strip(),
                    "",
                ]
            )
        )
    return "\n".join(blocks).strip()


def _refresh_canonical_transcript(manifest: dict[str, Any]) -> str:
    project_dir = run_dir(manifest["run_id"])
    canonical = _artifact_path(project_dir, "01_customer_discovery_conversation.md")
    canonical.parent.mkdir(parents=True, exist_ok=True)
    transcript_text = _combined_transcript_text(manifest)
    records = _transcript_records(manifest)
    content = [
        f"# {manifest['customer_name']} - Discovery Conversation\n\n",
        f"- Project: {manifest['project_name']}\n",
        f"- CSM: {manifest.get('csm_name') or 'Not captured'}\n",
        f"- Project ID: `{manifest['run_id']}`\n",
        f"- Canonical project root: `{project_dir}`\n",
        f"- Transcript sources: {len(records)}\n",
        f"- Last refreshed in cockpit: {utc_now()}\n\n",
        "## Transcript Corpus\n\n",
        transcript_text or "No transcript text has been attached yet.",
        "\n",
    ]
    canonical.write_text("".join(content), encoding="utf-8")
    manifest["transcript"] = {
        "source_path": records[-1].get("source_path", "") if records else "",
        "stored_path": records[-1].get("stored_path", "") if records else "",
        "canonical_path": str(canonical),
        "characters": len(transcript_text),
        "source_count": len(records),
        "analyzed_at": utc_now() if records else "",
        "demo_mode": any(record.get("is_demo") for record in records),
    }
    return transcript_text


def attach_transcript_from_path(
    manifest: dict[str, Any],
    source_path: Path,
    sections: list[Section],
    source_type: str = "local",
    original_name: str | None = None,
) -> dict[str, Any]:
    if not source_path.exists():
        raise FileNotFoundError(f"Transcript not found: {source_path}")

    text = read_transcript_text(source_path)
    target_dir = run_dir(manifest["run_id"]) / "data" / "raw" / "transcripts"
    target_dir.mkdir(parents=True, exist_ok=True)
    display_name = original_name or source_path.name
    target = source_path if source_path.parent.resolve() == target_dir.resolve() else _unique_transcript_target(target_dir, display_name)
    if source_path.resolve() != target.resolve():
        shutil.copy2(source_path, target)
    text_path = _write_transcript_text_file(target, text)
    is_demo = False
    try:
        is_demo = source_path.resolve() == DEFAULT_TRANSCRIPT_PATH.resolve()
    except FileNotFoundError:
        is_demo = False

    record = {
        "id": f"{datetime.now().strftime('%Y%m%d-%H%M%S')}-{slugify(Path(display_name).stem) or 'transcript'}",
        "source_type": source_type,
        "source_path": str(source_path),
        "original_name": display_name,
        "stored_path": str(target),
        "text_path": str(text_path),
        "characters": len(text),
        "attached_at": utc_now(),
        "is_demo": is_demo,
    }
    manifest.setdefault("transcripts", []).append(record)
    combined_text = _refresh_canonical_transcript(manifest)
    manifest["analysis"] = analyze_transcript_text(combined_text, sections, manifest.get("capture", {}))
    return manifest


def attach_uploaded_transcript(
    manifest: dict[str, Any],
    original_name: str,
    bytes_payload: bytes,
    sections: list[Section],
) -> dict[str, Any]:
    suffix = Path(original_name).suffix.lower()
    if suffix not in {".docx", ".md", ".txt"}:
        raise ValueError("Transcript upload must be .docx, .md, or .txt.")
    target_dir = run_dir(manifest["run_id"]) / "data" / "raw" / "transcripts"
    target_dir.mkdir(parents=True, exist_ok=True)
    target = _unique_transcript_target(target_dir, original_name)
    target.write_bytes(bytes_payload)
    return attach_transcript_from_path(manifest, target, sections, source_type="upload", original_name=original_name)


def calculate_readiness(manifest: dict[str, Any], sections: list[Section]) -> dict[str, Any]:
    capture = manifest.get("capture", {})
    required_sections = [section for section in sections if section.id in REQUIRED_FOR_SOP_GATE]

    capture_weights = {
        "answered": 1.0,
        "partial": 0.65,
        "needs_follow_up": 0.35,
        "not_answered": 0.0,
        "not_set": 0.0,
    }
    def avg(values: list[float]) -> int:
        return round((sum(values) / len(values)) * 100) if values else 0

    capture_pct = avg([capture_weights.get(capture.get(section.id, {}).get("status", "not_set"), 0) for section in required_sections])
    approval_pct = avg([1.0 if capture.get(section.id, {}).get("approved") else 0.0 for section in required_sections])
    doc_pct = avg([1.0 if manifest.get("artifacts", {}).get(name, {}).get("exists") else 0.0 for name in DOC_ARTIFACTS])
    overall_pct = round(capture_pct * 0.45 + approval_pct * 0.40 + doc_pct * 0.15)

    blockers: list[str] = []
    blocker_items: list[dict[str, str]] = []
    generation_missing: list[str] = []
    for section in required_sections:
        item = capture.get(section.id, {})
        capture_status = item.get("status", "not_set")
        if capture_status == "not_set":
            message = f"{section.label}: status not selected."
            blockers.append(message)
            blocker_items.append({"message": message, "section_id": section.id, "kind": "capture"})
            generation_missing.append(f"{section.label}: choose a status")
        elif capture_status in {"not_answered", "needs_follow_up"}:
            message = f"{section.label}: capture is {CAPTURE_STATUSES.get(capture_status, 'Not answered')}."
            blockers.append(message)
            blocker_items.append({"message": message, "section_id": section.id, "kind": "capture"})
        if not item.get("approved"):
            message = f"{section.label}: CSM approval is pending."
            blockers.append(message)
            blocker_items.append({"message": message, "section_id": section.id, "kind": "approval"})
            generation_missing.append(f"{section.label}: approve section")

    missing_docs = [ARTIFACT_LABELS[name] for name in DOC_ARTIFACTS if not manifest.get("artifacts", {}).get(name, {}).get("exists")]
    if missing_docs:
        message = "Generated document chain is incomplete."
        blockers.append(message)
        blocker_items.append({"message": message, "section_id": "", "kind": "documents"})

    sop_exists = manifest.get("artifacts", {}).get("03_accelerator_sop.md", {}).get("exists", False)
    generation_ready = not generation_missing
    return {
        "capture_pct": capture_pct,
        "approval_pct": approval_pct,
        "artifact_pct": doc_pct,
        "overall_pct": overall_pct,
        "workflow_gate": "ready" if not blockers and sop_exists else "blocked",
        "blockers": blockers[:14],
        "blocker_items": blocker_items[:14],
        "generation_ready": generation_ready,
        "generation_missing": generation_missing[:8],
        "required_section_count": len(required_sections),
    }


def _section_capture_md(section: Section, manifest: dict[str, Any]) -> str:
    item = manifest.get("capture", {}).get(section.id, {})
    analysis = manifest.get("analysis", {}).get(section.id, {})
    evidence = analysis.get("evidence", [])
    lines = [
        f"### {section.label}",
        "",
        f"- CSM status: {CAPTURE_STATUSES.get(item.get('status', 'not_set'), 'Not selected')}",
        f"- CSM approval: {'approved' if item.get('approved') else 'pending'}",
        f"- Transcript evidence: {EVIDENCE_STATUSES.get(analysis.get('status', 'not_run'), 'Not run')}",
        f"- Evidence summary: {analysis.get('summary', 'Transcript analysis has not run yet.')}",
        f"- Notes: {item.get('notes') or 'Not captured yet.'}",
    ]
    if evidence and analysis.get("status") != "supported":
        lines.append("- Review snippets:")
        for evidence_item in evidence:
            snippet = evidence_item["snippet"].replace("\n", " ")
            lines.append(f"  - {snippet}")
    lines.append("")
    return "\n".join(lines)


def _confirmed_sections(manifest: dict[str, Any], sections: list[Section]) -> list[Section]:
    return [
        section
        for section in sections
        if manifest.get("capture", {}).get(section.id, {}).get("approved")
        and manifest.get("capture", {}).get(section.id, {}).get("status") in {"answered", "partial"}
    ]


def _gap_rows(manifest: dict[str, Any], sections: list[Section]) -> list[dict[str, str]]:
    rows = []
    for section in sections:
        capture_item = manifest.get("capture", {}).get(section.id, {})
        analysis_item = manifest.get("analysis", {}).get(section.id, {})
        capture_status = capture_item.get("status", "not_set")
        evidence_status = analysis_item.get("status", "not_run")
        is_required = section.id in REQUIRED_FOR_SOP_GATE
        if (
            capture_status in {"not_set", "not_answered", "needs_follow_up"}
            or (is_required and not capture_item.get("approved"))
        ):
            rows.append(
                {
                    "section": section.label,
                    "required": "yes" if is_required else "supporting",
                    "capture": CAPTURE_STATUSES.get(capture_status, "Not selected"),
                    "evidence": EVIDENCE_STATUSES.get(evidence_status, "Not run"),
                    "approval": "approved" if capture_item.get("approved") else "pending",
                    "next_step": analysis_item.get("recommendation", "CSM follow-up required."),
                }
            )
    return rows


def _capture_note(manifest: dict[str, Any], section_id: str, fallback: str) -> str:
    note = manifest.get("capture", {}).get(section_id, {}).get("notes", "").strip()
    if note:
        return note
    analysis = manifest.get("analysis", {}).get(section_id, {})
    summary = analysis.get("summary", "").strip()
    if summary and analysis.get("status") in {"supported", "weak_evidence"}:
        return summary
    return fallback


def _missing_asset_inputs(manifest: dict[str, Any]) -> list[str]:
    missing = []
    required = {
        "business_problem": "Business problem",
        "current_process": "Current process",
        "desired_outcome": "Desired outcome",
        "value_realization": "Value realisation",
        "business_questions": "Business questions",
        "scope": "Scope",
        "inputs_sources_ownership": "Inputs, sources, and ownership",
        "rules_logic_definitions": "Rules, logic, and definitions",
        "exceptions_safe_handling": "Exceptions and safe handling",
        "validation_trust": "Validation/trust",
        "operational_readiness_phasing": "Operational readiness and phasing",
    }
    for section_id, label in required.items():
        item = manifest.get("capture", {}).get(section_id, {})
        if not item.get("notes", "").strip():
            missing.append(label)
    return missing


def _asset_gap_lines(manifest: dict[str, Any]) -> str:
    missing = _missing_asset_inputs(manifest)
    if not missing:
        return "- No customer-facing asset gaps detected from the approved capture.\n"
    return "".join(f"- {item}: add a clearer CSM/customer answer before customer reuse.\n" for item in missing)


def _asset_content_by_name(manifest: dict[str, Any], common_header: str) -> dict[str, str]:
    business_problem = _capture_note(manifest, "business_problem", "- Placeholder: business problem requires confirmation.")
    value_realization = _capture_note(manifest, "value_realization", "- Placeholder: value driver, baseline, target KPI, and measurement method require confirmation.")
    current_process = _capture_note(manifest, "current_process", "- Placeholder: current process requires confirmation.")
    desired_outcome = _capture_note(manifest, "desired_outcome", "- Placeholder: desired outcome requires confirmation.")
    business_questions = _capture_note(manifest, "business_questions", "- Placeholder: business questions require confirmation.")
    scope = _capture_note(manifest, "scope", "- Placeholder: first-slice scope and exclusions require confirmation.")
    inputs_sources = _capture_note(manifest, "inputs_sources_ownership", "- Placeholder: inputs, sources, ownership, known structure, and quality/access risks require confirmation.")
    rules_logic = _capture_note(manifest, "rules_logic_definitions", "- Placeholder: trusted rules, definitions, thresholds, and edge cases require confirmation.")
    exceptions_safe_handling = _capture_note(manifest, "exceptions_safe_handling", "- Placeholder: safe handling for missing, inconsistent, late, invalid, or risky records requires confirmation.")
    validation_trust = _capture_note(manifest, "validation_trust", "- Placeholder: validation method and sign-off owner require confirmation.")
    operational_readiness = _capture_note(manifest, "operational_readiness_phasing", "- Placeholder: operational readiness and phasing expectations require confirmation.")
    asset_gaps = _asset_gap_lines(manifest)

    value_statement = [
        common_header,
        "## Value Statement\n\n",
        "This is an upfront value hypothesis generated from discovery. It should be refined with the value team and revisited at close-out.\n\n",
        "## Reason For Doing This\n\n",
        business_problem,
        "\n\n## Value Hypothesis\n\n",
        value_realization,
        "\n\n## Baseline And Target\n\n",
        "- Baseline: captured in the Value Realisation notes, or still TBD.\n",
        "- Target KPI: captured in the Value Realisation notes, or still TBD.\n",
        "- Measurement method: captured in the Value Realisation notes, or still TBD.\n\n",
        "## Qualitative Value\n\n",
        desired_outcome,
        "\n\n## Assumptions And Gaps\n\n",
        asset_gaps,
        "\n## Close-Out Capture Placeholder\n\n",
        "- Final measured value: TBD after implementation.\n",
        "- Evidence source: TBD after implementation.\n",
        "- Customer quote or playback proof: TBD after implementation.\n",
    ]

    use_case_summary = [
        common_header,
        "## Use-Case Summary\n\n",
        "## Customer Wants To\n\n",
        desired_outcome,
        "\n\n## Current Pain\n\n",
        business_problem,
        "\n\n## Current Process\n\n",
        current_process,
        "\n\n## Business Questions\n\n",
        business_questions,
        "\n\n## Input Data\n\n",
        inputs_sources,
        "\n\n## Workflow Concept\n\n",
        "- Ingest the approved source extracts.\n",
        "- Normalize them into reusable accelerator entities.\n",
        "- Apply the captured business rules and validation checks.\n",
        "- Publish action-ready outputs for the intended users.\n\n",
        "## Outputs And Actions\n\n",
        desired_outcome,
        "\n\n## Expected Value\n\n",
        value_realization,
        "\n\n## Open Gaps\n\n",
        asset_gaps,
    ]

    case_study = [
        common_header,
        "## Case-Study Skeleton\n\n",
        "**Status:** pre-delivery draft. This is intentionally generated before implementation so the final story can be refined rather than reconstructed later.\n\n",
        "## Problem\n\n",
        business_problem,
        "\n\n## Solution Approach\n\n",
        desired_outcome,
        "\n\n",
        rules_logic,
        "\n\n## Expected Outcome\n\n",
        value_realization,
        "\n\n## Implementation Evidence To Add Later\n\n",
        "- Final workflow/package link: TBD.\n",
        "- Final measured impact: TBD.\n",
        "- Customer validation quote: TBD.\n",
        "- Reusable accelerator asset location: TBD.\n\n",
        "## Reuse Notes\n\n",
        scope,
        "\n",
    ]

    accelerator_101 = [
        common_header,
        "## Accelerator 101 - Generic Guide\n\n",
        "## Business Problem\n\n",
        business_problem,
        "\n\n## Prerequisites\n\n",
        "- A named business owner and CSM/consultant owner.\n",
        "- Approved first-slice scope.\n",
        "- Sample-safe input extracts or synthetic equivalents.\n\n",
        "## Input Contract\n\n",
        inputs_sources,
        "\n\n",
        exceptions_safe_handling,
        "\n\n## Modular Workflow Architecture\n\n",
        "- Source contract and profiling.\n",
        "- Neutral schema normalization.\n",
        "- Business rule and exception layer.\n",
        "- Prioritization or action logic.\n",
        "- Output publishing and governance review.\n\n",
        "## Step-By-Step Usage\n\n",
        "- Confirm the use case and first-slice scope.\n",
        "- Map customer inputs to the generic input contract.\n",
        "- Run the starter workflow against sample data.\n",
        "- Review outputs with the business owner.\n",
        "- Capture gaps for 102/201 deepening.\n\n",
        "## KPIs\n\n",
        value_realization,
        "\n\n## Limitations And Reuse Guidance\n\n",
        asset_gaps,
    ]

    accelerator_102 = [
        common_header,
        "## Accelerator 102 - Starter Kit Guide\n\n",
        "## Starter Kit Shape\n\n",
        "- 101 generic guide.\n",
        "- 102 starter-kit notes.\n",
        "- 201 implementation note when the source/system pattern is known.\n",
        "- Sample or synthetic inputs.\n",
        "- Workflow output and validation expectations.\n\n",
        "## Configuration Choices\n\n",
        scope,
        "\n\n## Sample Data Expectations\n\n",
        inputs_sources,
        "\n\n",
        exceptions_safe_handling,
        "\n\n## Validation Plan\n\n",
        validation_trust,
        "\n\n## Adaptation Notes\n\n",
        "- Keep the customer-specific fields mapped to neutral entities.\n",
        "- Keep output names and KPIs aligned with the approved value statement.\n",
        "- Do not promote assumptions to facts without CSM approval.\n\n",
        "## Gaps To Resolve Before Reuse\n\n",
        asset_gaps,
    ]

    accelerator_201 = [
        common_header,
        "## Accelerator 201 - Implementation Note\n\n",
        "## Source Context\n\n",
        inputs_sources,
        "\n\n## Field And Entity Mapping\n\n",
        inputs_sources,
        "\n\n## Module-Level Build Logic\n\n",
        "- Input module: validate source contract, required fields, and grain.\n",
        "- Normalization module: map customer-shaped inputs into reusable accelerator entities.\n",
        "- Rules module: apply approved definitions, thresholds, and exception logic.\n",
        "- Output module: publish action, summary, and governance views.\n\n",
        "## Business Rules\n\n",
        rules_logic,
        "\n\n## Outputs\n\n",
        desired_outcome,
        "\n\n## Governance And Monitoring\n\n",
        validation_trust,
        "\n\n",
        operational_readiness,
        "\n\n## System-Specific Assumptions\n\n",
        asset_gaps,
    ]

    return {
        "04_value_statement.md": "".join(value_statement),
        "05_use_case_summary.md": "".join(use_case_summary),
        "06_case_study_skeleton.md": "".join(case_study),
        "07_accelerator_101.md": "".join(accelerator_101),
        "08_accelerator_102.md": "".join(accelerator_102),
        "09_accelerator_201.md": "".join(accelerator_201),
    }


def _relative_to_project(project_dir: Path, path: Path) -> str:
    try:
        return path.relative_to(project_dir).as_posix()
    except ValueError:
        return os.path.relpath(path, project_dir).replace("\\", "/")


def _first_existing(paths: list[Path]) -> Path | None:
    for path in paths:
        if path.exists():
            return path
    return None


def _detect_codex_command() -> str:
    """Return a launchable Codex command path when the local app exposes one.

    This does not detect an already-open Codex Desktop conversation. It only
    detects whether the cockpit process can launch a local Codex command.
    """
    for env_name in ("CODEX_PATH", "CODEX_CLI", "OPENAI_CODEX_PATH"):
        raw_path = os.environ.get(env_name, "").strip().strip('"')
        if raw_path and Path(raw_path).exists():
            return str(Path(raw_path))

    for command in ("codex", "codex.exe"):
        detected = shutil.which(command)
        if detected:
            return detected

    if os.name == "nt":
        user_windows_apps = Path(os.environ.get("LOCALAPPDATA", "")) / "Microsoft" / "WindowsApps"
        alias = user_windows_apps / "codex.exe"
        if alias.exists():
            return str(alias)

        package_roots = [
            Path(os.environ.get("ProgramFiles", r"C:\Program Files")) / "WindowsApps",
            Path(os.environ.get("ProgramW6432", r"C:\Program Files")) / "WindowsApps",
        ]
        candidates: list[Path] = []
        for package_root in package_roots:
            if not package_root.exists():
                continue
            candidates.extend(package_root.glob("OpenAI.Codex_*_*__*/*/resources/codex.exe"))
            candidates.extend(package_root.glob("OpenAI.Codex_*_*__*/app/resources/codex.exe"))
        existing = [candidate for candidate in candidates if candidate.exists()]
        if existing:
            existing.sort(key=lambda path: path.stat().st_mtime, reverse=True)
            return str(existing[0])

    return ""


def detect_workflow_environment() -> dict[str, Any]:
    program_files = Path(os.environ.get("ProgramFiles", r"C:\Program Files"))
    program_files_x86 = Path(os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)"))
    codex_path = _detect_codex_command()
    engine = _first_existing(
        [
            program_files / "Alteryx" / "bin" / "AlteryxEngineCmd.exe",
            program_files_x86 / "Alteryx" / "bin" / "AlteryxEngineCmd.exe",
        ]
    )
    designer = _first_existing(
        [
            program_files / "Alteryx" / "bin" / "AlteryxGui.exe",
            program_files / "Alteryx" / "bin" / "AlteryxDesigner.exe",
            program_files / "Alteryx" / "bin" / "AlteryxDesigner_x64.exe",
            program_files_x86 / "Alteryx" / "bin" / "AlteryxGui.exe",
            program_files_x86 / "Alteryx" / "bin" / "AlteryxDesigner.exe",
        ]
    )
    builder_root = TOOLING_DIR / "alteryx_workflow_builder"
    beautification_root = TOOLING_DIR / "alteryx-beautification"
    hybrid_builder_reference = builder_root / "golden" / "customer_facing_hybrid_reference" / "10_REFERENCE_customer_facing_hybrid.yxmd"
    hybrid_beauty_profile = beautification_root / "references" / "customer-facing-hybrid-profile.md"
    hybrid_beauty_reference = beautification_root / "assets" / "customer-facing-hybrid-reference.yxmd"
    builder_ready = (
        (builder_root / "verify_workflows.py").exists()
        and (builder_root / "WORKFLOW_RULES.md").exists()
        and hybrid_builder_reference.exists()
    )
    beautification_ready = (
        (beautification_root / "SKILL.md").exists()
        and (beautification_root / "references" / "spiderweb-reduction.md").exists()
        and hybrid_beauty_profile.exists()
        and hybrid_beauty_reference.exists()
    )
    return {
        "codex_detected": bool(codex_path),
        "codex_path": codex_path,
        "codex_detection_note": "Launch command available" if codex_path else "No launchable Codex command found; an already-open Codex chat cannot be detected from the local web app.",
        "designer_detected": bool(engine or designer),
        "engine_path": str(engine) if engine else "",
        "designer_path": str(designer) if designer else "",
        "builder_toolkit_ready": builder_ready,
        "beautification_ready": beautification_ready,
        "tooling_bundle_version": f"cockpit-v{APP_VERSION}",
        "builder_toolkit_path": str(builder_root),
        "beautification_path": str(beautification_root),
    }


def detect_workflow_artifacts(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    project_dir = run_dir(manifest["run_id"])
    candidates: list[Path] = []
    workflow_dirs = [project_dir / WORKFLOW_OUTPUT_DIR, project_dir / "workflows"]
    validation_dirs = [project_dir / WORKFLOW_VALIDATION_DIR, project_dir / "validation"]
    status_dirs = [project_dir / WORKFLOW_STATUS_DIR, project_dir / "status"]
    for workflow_dir in workflow_dirs:
        if workflow_dir.exists():
            for pattern in ["*.yxmd", "*.yxmc", "*.yxwz", "*.yxzp", "*.png", "*.json", "*.md", "*.csv", "*.xlsx"]:
                candidates.extend(workflow_dir.rglob(pattern))
    for validation_dir in validation_dirs:
        if validation_dir.exists():
            for pattern in ["*.png", "*.json", "*.md", "*.csv", "*.xlsx", "*.txt"]:
                candidates.extend(validation_dir.rglob(pattern))
    for status_dir in status_dirs:
        for status_name in ["workflow_spec.json", "validation_report.json", "lint_report.json", "runtime_smoke_results.json"]:
            status_path = status_dir / status_name
            if status_path.exists():
                candidates.append(status_path)

    seen: set[Path] = set()
    artifacts: list[dict[str, Any]] = []
    for path in sorted(candidates):
        if path in seen or not path.is_file():
            continue
        seen.add(path)
        artifacts.append(
            {
                "name": path.name,
                "relative_path": _relative_to_project(project_dir, path),
                "kind": path.suffix.lower().lstrip(".") or "file",
                "size": path.stat().st_size,
                "updated_at": datetime.fromtimestamp(path.stat().st_mtime, timezone.utc).replace(microsecond=0).isoformat(),
            }
        )
    return artifacts


def refresh_workflow_build_state(manifest: dict[str, Any]) -> dict[str, Any]:
    project_dir = run_dir(manifest["run_id"])
    preflight = detect_workflow_environment()
    prompt_path = project_dir / WORKFLOW_PROMPT_ARTIFACT
    helper_path = project_dir / WORKFLOW_HELPER_ARTIFACT
    workflow_manifest_path = project_dir / WORKFLOW_MANIFEST_ARTIFACT
    detected_artifacts = detect_workflow_artifacts(manifest)
    workflow_files = [
        item
        for item in detected_artifacts
        if item["relative_path"].startswith("workflows/") and item["kind"] in {"yxmd", "yxmc", "yxwz", "yxzp"}
    ]
    existing = manifest.setdefault("workflow_build", {})
    generated_at = existing.get("generated_at", "")
    if workflow_files:
        status = "workflow_outputs_detected"
    elif prompt_path.exists():
        status = "prompt_ready"
    else:
        status = existing.get("status", "not_generated")
        if status == "blocked":
            status = "not_generated"

    existing.update(
        {
            "status": status,
            "prompt_path": str(prompt_path) if prompt_path.exists() else "",
            "helper_script_path": str(helper_path) if helper_path.exists() else "",
            "manifest_path": str(workflow_manifest_path) if workflow_manifest_path.exists() else "",
            "prompt_artifact": WORKFLOW_PROMPT_ARTIFACT,
            "helper_artifact": WORKFLOW_HELPER_ARTIFACT,
            "manifest_artifact": WORKFLOW_MANIFEST_ARTIFACT,
            "codex_detected": preflight["codex_detected"],
            "codex_path": preflight["codex_path"],
            "codex_detection_note": preflight["codex_detection_note"],
            "designer_detected": preflight["designer_detected"],
            "engine_path": preflight["engine_path"],
            "designer_path": preflight["designer_path"],
            "tooling_bundle_version": preflight["tooling_bundle_version"],
            "builder_toolkit_ready": preflight["builder_toolkit_ready"],
            "beautification_ready": preflight["beautification_ready"],
            "detected_artifacts": detected_artifacts[:40],
        }
    )
    if generated_at:
        existing["generated_at"] = generated_at
    return manifest


def _pipeline_status(manifest: dict[str, Any], readiness: dict[str, Any]) -> dict[str, Any]:
    sequence = load_sequence_config()
    workflow_build = manifest.get("workflow_build", {})
    peer_review = manifest.get("peer_review", {})
    return {
        "run_id": manifest["run_id"],
        "customer_name": manifest["customer_name"],
        "project_name": manifest["project_name"],
        "updated_at": utc_now(),
        "readiness": readiness,
        "workflow_build": {
            "state": readiness["workflow_gate"],
            "large_prompt_only": True,
            "automatic_generation": False,
            "handoff_status": workflow_build.get("status", "not_generated"),
            "codex_detected": workflow_build.get("codex_detected", False),
            "designer_detected": workflow_build.get("designer_detected", False),
            "builder_toolkit_ready": workflow_build.get("builder_toolkit_ready", False),
            "beautification_ready": workflow_build.get("beautification_ready", False),
        },
        "peer_review": peer_review,
        "stages": sequence.get("stages", []),
    }


def _write_peer_review_status(manifest: dict[str, Any], readiness: dict[str, Any]) -> dict[str, Any]:
    project_dir = run_dir(manifest["run_id"])
    status_path = _artifact_path(project_dir, PEER_REVIEW_STATUS_ARTIFACT)
    status_path.parent.mkdir(parents=True, exist_ok=True)
    existing_state = manifest.get("peer_review", {}).get("state", "draft")
    all_assets_exist = all(_artifact_path(project_dir, artifact).exists() for artifact in ACCELERATOR_ASSET_ARTIFACTS)
    if existing_state in {"reviewed", "published"}:
        state = existing_state
    elif readiness["workflow_gate"] == "ready" and all_assets_exist:
        state = "ready_for_peer_review"
    else:
        state = "draft"
    peer_review = {
        "state": state,
        "updated_at": utc_now(),
        "customer_facing_assets": ACCELERATOR_ASSET_ARTIFACTS,
        "ready_for_peer_review": state in {"ready_for_peer_review", "reviewed", "published"},
        "notes": "Customer-facing accelerator assets are generated as Markdown drafts for CSM/Jon review.",
        "blockers": [] if state != "draft" else [
            "SOP gate must be ready and all accelerator assets must exist before peer review."
        ],
    }
    status_path.write_text(json.dumps(peer_review, indent=2), encoding="utf-8")
    manifest["peer_review"] = peer_review
    manifest.setdefault("artifacts", {}).setdefault(PEER_REVIEW_STATUS_ARTIFACT, {})
    manifest["artifacts"][PEER_REVIEW_STATUS_ARTIFACT]["generated_at"] = peer_review["updated_at"]
    return manifest


def _sanitize_large_prompt_for_handoff(large_prompt: str) -> str:
    sanitized = large_prompt
    sanitized = re.sub(
        r"Suggested usage pattern:\s*```text.*?```\s*",
        "Suggested usage pattern:\n\nThe project-specific usage pattern has been resolved in the Case Inputs section above.\n\n",
        sanitized,
        flags=re.DOTALL,
    )
    sanitized = re.sub(
        r"Required full build stack:\s*1\..*?If this stack is not followed, the build is out of process\.\s*",
        (
            "Required full build stack:\n\n"
            "Use the generated Case Inputs section above as the resolved build stack for this project. "
            "Do not reintroduce generic placeholders.\n\n"
        ),
        sanitized,
        flags=re.DOTALL,
    )
    sanitized = sanitized.replace("[INSERT PATH]", "resolved in Case Inputs above")
    sanitized = sanitized.replace("[INSERT SOURCE STYLE]", "resolved from the approved source-system notes")
    return sanitized


def _workflow_handoff_prompt(manifest: dict[str, Any], readiness: dict[str, Any], preflight: dict[str, Any]) -> str:
    large_prompt = LARGE_PROMPT_PATH.read_text(encoding="utf-8") if LARGE_PROMPT_PATH.exists() else "Large prompt pack is missing from the process pack."
    large_prompt = _sanitize_large_prompt_for_handoff(large_prompt)
    project_dir = run_dir(manifest["run_id"])
    identity = manifest.get("project_identity", {})
    identity_hash = identity.get("identity_hash", "")
    state = "READY" if readiness["workflow_gate"] == "ready" else "BLOCKED"
    blockers = "\n".join(f"- {blocker}" for blocker in readiness["blockers"]) or "- None"
    builder_path = project_dir / "tooling" / "alteryx_workflow_builder"
    beautification_path = project_dir / "tooling" / "alteryx-beautification"
    hybrid_builder_reference = builder_path / "golden" / "customer_facing_hybrid_reference" / "10_REFERENCE_customer_facing_hybrid.yxmd"
    hybrid_beauty_profile = beautification_path / "references" / "customer-facing-hybrid-profile.md"
    hybrid_beauty_reference = beautification_path / "assets" / "customer-facing-hybrid-reference.yxmd"
    transcript_count = manifest.get("transcript", {}).get("source_count", len(_transcript_records(manifest)))
    case_inputs = {
        "Canonical project root": str(project_dir),
        "Project ID": manifest["run_id"],
        "Project identity hash": identity_hash,
        "Project manifest": str(project_dir / "manifest.json"),
        "Discovery conversation": _relative_to_project(project_dir, _artifact_path(project_dir, "01_customer_discovery_conversation.md")),
        "Guided SOP capture": _relative_to_project(project_dir, _artifact_path(project_dir, "02_guided_sop_capture.md")),
        "Accelerator SOP": _relative_to_project(project_dir, _artifact_path(project_dir, "03_accelerator_sop.md")),
        "Gap log": _relative_to_project(project_dir, _artifact_path(project_dir, "sop_gap_log.md")),
        "Architecture assessment": _relative_to_project(project_dir, _artifact_path(project_dir, "sop_architecture_assessment.md")),
        "Value statement": _relative_to_project(project_dir, _artifact_path(project_dir, "04_value_statement.md")),
        "Use-case summary": _relative_to_project(project_dir, _artifact_path(project_dir, "05_use_case_summary.md")),
        "Case-study skeleton": _relative_to_project(project_dir, _artifact_path(project_dir, "06_case_study_skeleton.md")),
        "Accelerator 101": _relative_to_project(project_dir, _artifact_path(project_dir, "07_accelerator_101.md")),
        "Accelerator 102": _relative_to_project(project_dir, _artifact_path(project_dir, "08_accelerator_102.md")),
        "Accelerator 201": _relative_to_project(project_dir, _artifact_path(project_dir, "09_accelerator_201.md")),
        "Peer-review status": _relative_to_project(project_dir, _artifact_path(project_dir, PEER_REVIEW_STATUS_ARTIFACT)),
        "Raw transcript folder": "data/raw/transcripts",
        "Generated data folder": "data/generated",
        "Workflow output folder": WORKFLOW_OUTPUT_DIR,
        "Validation output folder": WORKFLOW_VALIDATION_DIR,
        "Workflow build manifest": _relative_to_project(project_dir, _artifact_path(project_dir, WORKFLOW_MANIFEST_ARTIFACT)),
        "Alteryx workflow-builder toolkit": _relative_to_project(project_dir, builder_path),
        "Hybrid workflow reference": _relative_to_project(project_dir, hybrid_builder_reference),
        "Beautification rules": _relative_to_project(project_dir, beautification_path),
        "Hybrid beautification profile": _relative_to_project(project_dir, hybrid_beauty_profile),
        "Hybrid beautification reference": _relative_to_project(project_dir, hybrid_beauty_reference),
        "Transcript source count": str(transcript_count),
        "Reusable Large prompt source": str(LARGE_PROMPT_PATH),
    }
    case_input_lines = "\n".join(f"- {label}: `{value}`" for label, value in case_inputs.items())
    detected = "\n".join(
        [
            f"- Codex launch command: {'available' if preflight['codex_detected'] else 'not found'}",
            f"- Alteryx Designer/Engine detected: {'yes' if preflight['designer_detected'] else 'no'}",
            f"- Builder toolkit bundled: {'yes' if preflight['builder_toolkit_ready'] else 'no'}",
            f"- Beautification guidance bundled: {'yes' if preflight['beautification_ready'] else 'no'}",
        ]
    )
    return f"""# Codex Workflow Build Prompt: {manifest['customer_name']} / {manifest['project_name']}

State: {state}
Generated: {utc_now()}
Project ID: `{manifest['run_id']}`
Canonical project root: `{project_dir}`
Project identity hash: `{identity_hash}`

This is the project-specific, hydrated workflow-build prompt. Do not edit Jon's reusable Large prompt source. If the reusable prompt changes in Git, regenerate this handoff from the cockpit so the project keeps the latest approved process with the correct case inputs.

## Mandatory Identity And Path Gate

Before doing any workflow design or file search, prove that you are in the correct project.

1. Resolve the current working directory and compare it to the canonical project root above.
2. Read `manifest.json` from the canonical project root and confirm `run_id` equals `{manifest['run_id']}`.
3. Read `{ARTIFACT_PATHS[WORKFLOW_MANIFEST_ARTIFACT]}` and confirm `project_identity_hash` equals `{identity_hash}`.
4. Confirm `01_discovery/01_customer_discovery_conversation.md`, `02_sop_authoring/02_guided_sop_capture.md`, `02_sop_authoring/03_accelerator_sop.md`, and `{ARTIFACT_PATHS[WORKFLOW_MANIFEST_ARTIFACT]}` exist under the canonical project root.
5. If any identity, path, or SOP-gate check fails, stop. Do not search parent folders, sibling projects, demo folders, shelf projects, Downloads, or customer folders as fallbacks.

## Operating Mode

You are Codex running locally with access to this project folder. Build the Alteryx Designer workflow only after checking the SOP gate state below.

- If State is `BLOCKED`, do not build the workflow. Read the gap log, explain what is missing, and update the project plan only if asked.
- If State is `READY`, build the workflow from the approved SOP and case inputs below.
- Keep all generated assets inside this project folder. Do not reference customer source folders or external systems.
- Use sanitized/sample data only unless the user explicitly provides approved local inputs.
- Treat bundled demo transcripts as demo/training context only. Never substitute another nearby accelerator project when this project's files are missing.
- Treat the Jon folder chain as the workflow source of truth. The transcript corpus is raw evidence; do not bypass the approved SOP/doc chain and build directly from transcript text.
- Use the value statement, use-case summary, case-study skeleton, and 101/102/201 drafts as supporting context. They should improve narrative alignment, but customer-facing polish must not block the workflow build if the SOP gate is ready.

## Case Inputs

{case_input_lines}

## Current Blockers

{blockers}

## Local Preflight

{detected}

## Mandatory Beautification Layer

Apply the bundled Alteryx beautification rules as build requirements, not as optional polish.

- Use a title-first canvas with a clear left-to-right narrative.
- Place every real tool inside a contextual container with compact, useful annotations.
- Use the bundled customer-facing hybrid reference workflow and profile as the visual standard: top title/banner, clean tool lane, bottom documentation shelf, contextual colored containers, and minimal visual clutter.
- Apply aggressive spiderweb reduction: minimize connector crossings, avoid crowded fan-in knots, separate branch lanes, and reroute geometry until the rendered workflow reads cleanly.
- Prefer readability over blind tool-count minimization. Condense only where it improves both maintainability and visual scanability.
- Render the workflow preview and iterate until the canvas is clean, balanced, readable, and package-safe.

## Expected Build Outputs

- `{WORKFLOW_OUTPUT_DIR}/main.yxmd` or a clearly named first-slice `.yxmd`.
- Supporting macros only if they materially improve reuse.
- Synthetic/sample input data under `data/generated`.
- Validation outputs under `{WORKFLOW_VALIDATION_DIR}`.
- `{WORKFLOW_STATUS_DIR}/workflow_spec.json` describing the built workflow.
- Rendered preview images proving the workflow is visually reviewable.
- Updated gap/architecture notes if the SOP was insufficient for implementation.

## Launch Note

If this prompt was opened by the cockpit, it may also have been copied to the clipboard. Paste the entire prompt into local Codex from the project root, then let Codex use the files and tooling paths above.

---

{large_prompt}
"""


def _workflow_helper_script(manifest: dict[str, Any]) -> str:
    project_dir = run_dir(manifest["run_id"])
    identity_hash = manifest.get("project_identity", {}).get("identity_hash", "")
    run_id = manifest["run_id"]
    project_dir_literal = str(project_dir).replace("'", "''")
    script = """$ErrorActionPreference = "Stop"

$ExpectedProjectRoot = @'
__PROJECT_ROOT__
'@
$ProjectRoot = (Resolve-Path -LiteralPath $ExpectedProjectRoot).Path
$StatusDir = Join-Path $ProjectRoot "03_workflow_build\\status"
$PromptPath = Join-Path $StatusDir "codex_workflow_build_prompt.md"
$BuildManifestPath = Join-Path $StatusDir "workflow_build_manifest.json"
$ProjectManifestPath = Join-Path $ProjectRoot "manifest.json"

if (-not (Test-Path $PromptPath)) {
    Write-Host "Prompt not found: $PromptPath"
    Read-Host "Press Enter to close"
    exit 1
}
if (-not (Test-Path $ProjectManifestPath) -or -not (Test-Path $BuildManifestPath)) {
    Write-Host "Project identity files are missing. Refusing to launch Codex against an unsafe folder."
    Read-Host "Press Enter to close"
    exit 1
}

$ProjectManifest = Get-Content -Raw -LiteralPath $ProjectManifestPath | ConvertFrom-Json
$BuildManifest = Get-Content -Raw -LiteralPath $BuildManifestPath | ConvertFrom-Json
if ($ProjectManifest.run_id -ne "__RUN_ID__" -or $BuildManifest.project_identity_hash -ne "__IDENTITY_HASH__") {
    Write-Host "Project identity mismatch. Refusing fallback search or launch."
    Write-Host "Expected run id: __RUN_ID__"
    Write-Host "Expected identity: __IDENTITY_HASH__"
    Read-Host "Press Enter to close"
    exit 1
}

$promptText = Get-Content -Raw -LiteralPath $PromptPath
try {
    Set-Clipboard -Value $promptText
    Write-Host "Codex workflow prompt copied to clipboard."
} catch {
    Write-Host "Could not copy to clipboard. Open this file manually:"
    Write-Host $PromptPath
}

$codexCommand = Get-Command codex -ErrorAction SilentlyContinue
if ($codexCommand) {
    Write-Host "Opening Codex in project folder:"
    Write-Host $ProjectRoot
    Start-Process -FilePath $codexCommand.Source -WorkingDirectory $ProjectRoot
} else {
    Write-Host "Codex was not found on PATH. Open Codex manually in this folder:"
    Write-Host $ProjectRoot
}

Write-Host ""
Write-Host "Paste the clipboard prompt into Codex if it does not appear automatically."
Read-Host "Press Enter to close"
"""
    return script.replace("__PROJECT_ROOT__", project_dir_literal).replace("__RUN_ID__", run_id).replace("__IDENTITY_HASH__", identity_hash)


def generate_workflow_build_handoff(
    manifest: dict[str, Any],
    sections: list[Section],
    readiness: dict[str, Any] | None = None,
) -> dict[str, Any]:
    project_dir = run_dir(manifest["run_id"])
    ensure_project_structure(project_dir)
    status_dir = project_dir / WORKFLOW_STATUS_DIR
    status_dir.mkdir(parents=True, exist_ok=True)
    readiness = readiness or calculate_readiness(_refresh_artifact_records(manifest), sections)
    preflight = detect_workflow_environment()
    prompt_path = _artifact_path(project_dir, WORKFLOW_PROMPT_ARTIFACT)
    helper_path = _artifact_path(project_dir, WORKFLOW_HELPER_ARTIFACT)
    build_manifest_path = _artifact_path(project_dir, WORKFLOW_MANIFEST_ARTIFACT)
    prompt = _workflow_handoff_prompt(manifest, readiness, preflight)

    prompt_path.write_text(prompt, encoding="utf-8")
    _artifact_path(project_dir, LEGACY_PROMPT_ARTIFACT).write_text(prompt, encoding="utf-8")
    helper_path.write_text(_workflow_helper_script(manifest), encoding="utf-8")

    build_manifest = {
        "schema_version": 2,
        "generated_at": utc_now(),
        "run_id": manifest["run_id"],
        "canonical_project_root": str(project_dir),
        "project_identity_hash": manifest.get("project_identity", {}).get("identity_hash", ""),
        "customer_name": manifest["customer_name"],
        "project_name": manifest["project_name"],
        "state": readiness["workflow_gate"],
        "case_inputs": {
            "discovery": "01_discovery",
            "sop_authoring": "02_sop_authoring",
            "workflow_build": "03_workflow_build",
            "reference_examples": "04_reference_examples",
            "sequencer": "sequencer",
            "data_raw": "data/raw",
            "data_generated": "data/generated",
            "status": WORKFLOW_STATUS_DIR,
            "workflows": WORKFLOW_OUTPUT_DIR,
            "validation": WORKFLOW_VALIDATION_DIR,
            "prompt": WORKFLOW_PROMPT_ARTIFACT,
            "helper": WORKFLOW_HELPER_ARTIFACT,
            "project_manifest": "manifest.json",
            "accelerator_sop": ARTIFACT_PATHS["03_accelerator_sop.md"],
            "value_statement": ARTIFACT_PATHS["04_value_statement.md"],
            "use_case_summary": ARTIFACT_PATHS["05_use_case_summary.md"],
            "case_study_skeleton": ARTIFACT_PATHS["06_case_study_skeleton.md"],
            "accelerator_101": ARTIFACT_PATHS["07_accelerator_101.md"],
            "accelerator_102": ARTIFACT_PATHS["08_accelerator_102.md"],
            "accelerator_201": ARTIFACT_PATHS["09_accelerator_201.md"],
            "peer_review_status": ARTIFACT_PATHS[PEER_REVIEW_STATUS_ARTIFACT],
        },
        "absolute_case_inputs": {
            "project_root": str(project_dir),
            "project_manifest": str(project_dir / "manifest.json"),
            "workflow_build_manifest": str(build_manifest_path),
            "prompt": str(prompt_path),
            "discovery": str(project_dir / "01_discovery"),
            "sop_authoring": str(project_dir / "02_sop_authoring"),
            "workflow_build": str(project_dir / "03_workflow_build"),
            "workflows": str(project_dir / WORKFLOW_OUTPUT_DIR),
            "validation": str(project_dir / WORKFLOW_VALIDATION_DIR),
        },
        "transcripts": manifest.get("transcripts", []),
        "demo_mode": manifest.get("transcript", {}).get("demo_mode", False),
        "preflight": preflight,
        "expected_outputs": [
            f"{WORKFLOW_OUTPUT_DIR}/main.yxmd",
            f"{WORKFLOW_STATUS_DIR}/workflow_spec.json",
            f"{WORKFLOW_VALIDATION_DIR}/validation_report.json",
            f"{WORKFLOW_VALIDATION_DIR}/workflow_preview.png",
        ],
        "blockers": readiness["blockers"],
        "notes": "Codex is the workflow-building brain. The cockpit prepares context, prompts, tooling, and gates.",
    }
    build_manifest_path.write_text(json.dumps(build_manifest, indent=2), encoding="utf-8")

    generated_at = utc_now()
    for artifact in [WORKFLOW_PROMPT_ARTIFACT, WORKFLOW_HELPER_ARTIFACT, WORKFLOW_MANIFEST_ARTIFACT, LEGACY_PROMPT_ARTIFACT]:
        manifest.setdefault("artifacts", {}).setdefault(artifact, {})
        manifest["artifacts"][artifact]["generated_at"] = generated_at

    manifest.setdefault("workflow_build", {}).update(
        {
            "status": "prompt_ready" if readiness["workflow_gate"] == "ready" else "blocked",
            "prompt_path": str(prompt_path),
            "helper_script_path": str(helper_path),
            "manifest_path": str(build_manifest_path),
            "generated_at": generated_at,
            **{key: preflight[key] for key in ["codex_detected", "codex_path", "codex_detection_note", "designer_detected", "engine_path", "designer_path", "tooling_bundle_version", "builder_toolkit_ready", "beautification_ready"]},
        }
    )
    return _refresh_artifact_records(manifest)


def generate_docs(manifest: dict[str, Any], sections: list[Section]) -> dict[str, Any]:
    project_dir = run_dir(manifest["run_id"])
    ensure_project_structure(project_dir)
    status_dir = project_dir / WORKFLOW_STATUS_DIR
    readiness = calculate_readiness(_refresh_artifact_records(manifest), sections)
    confirmed = _confirmed_sections(manifest, sections)
    gap_rows = _gap_rows(manifest, sections)

    common_header = (
        f"# {manifest['customer_name']} - {manifest['project_name']}\n\n"
        f"- CSM: {manifest.get('csm_name') or 'Not captured'}\n"
        f"- Project ID: `{manifest['run_id']}`\n"
        f"- Canonical project root: `{run_dir(manifest['run_id'])}`\n"
        f"- Transcript sources: {manifest.get('transcript', {}).get('source_count', len(_transcript_records(manifest)))}\n"
        f"- Generated: {utc_now()}\n"
        f"- Process: Jon accelerator operating system, Large-only workflow handoff\n\n"
    )

    discovery = [
        common_header,
        "## Discovery Capture\n\n",
        "This file is the CSM-safe discovery handoff. It summarizes capture status, transcript support, assumptions, and gaps without requiring the next person to read a raw transcript first.\n\n",
    ]
    discovery.extend(_section_capture_md(section, manifest) for section in sections)

    guided = [common_header, "## Confirmed Facts\n\n"]
    if confirmed:
        for section in confirmed:
            notes = manifest["capture"][section.id].get("notes") or manifest.get("analysis", {}).get(section.id, {}).get("summary", "Confirmed by transcript evidence.")
            guided.append(f"- **{section.label}:** {notes}\n")
    else:
        guided.append("- No sections have been approved yet.\n")
    guided.append("\n## Assumptions And Open Questions\n\n")
    if gap_rows:
        for row in gap_rows:
            guided.append(f"- **{row['section']} ({row['required']}):** {row['next_step']}\n")
    else:
        guided.append("- No open gaps detected.\n")
    guided.append("\n## Candidate Workflow Modules\n\n")
    guided.append("- Input standardisation and source contract validation.\n")
    guided.append("- Data-shape normalization into neutral accelerator entities.\n")
    guided.append("- Business rule application and exception routing.\n")
    guided.append("- Prioritisation, scoring, and output publishing.\n")
    guided.append("- Governance review output for weak or unsafe cases.\n")

    sop = [
        common_header,
        "## Build Objective\n\n",
        "Create a first-slice accelerator workflow only after the SOP gate is satisfied. The workflow should reflect approved customer facts, explicit assumptions, and known open gaps.\n\n",
        "## Scope And Business Outcome\n\n",
        manifest.get("capture", {}).get("scope", {}).get("notes") or "- Scope still requires CSM approval.\n",
        "\n\n## Value Realisation\n\n",
        manifest.get("capture", {}).get("value_realization", {}).get("notes") or "- Value driver, baseline, and measurement approach still require approval.\n",
        "\n\n## Inputs, Sources, And Ownership\n\n",
        manifest.get("capture", {}).get("inputs_sources_ownership", {}).get("notes") or "- Inputs, ownership, source structure, quality, and access constraints still require approval.\n",
        "\n\n## Rules, Logic, And Definitions\n\n",
        manifest.get("capture", {}).get("rules_logic_definitions", {}).get("notes") or "- Business definitions, thresholds, filters, and exception rules still require approval.\n",
        "\n\n## Exceptions And Safe Handling\n\n",
        manifest.get("capture", {}).get("exceptions_safe_handling", {}).get("notes") or "- Safe handling for missing, invalid, late, or risky data still requires approval.\n",
        "\n\n## Outputs And Actions\n\n",
        manifest.get("capture", {}).get("desired_outcome", {}).get("notes") or "- Output/action design still requires approval.\n",
        "\n\n## Validation Expectations\n\n",
        manifest.get("capture", {}).get("validation_trust", {}).get("notes") or "- Validation method and sign-off owner still require approval.\n",
        "\n\n## Operational Readiness And Phasing\n\n",
        manifest.get("capture", {}).get("operational_readiness_phasing", {}).get("notes") or "- Phase 1 operations and later-state production needs still require approval.\n",
        "\n\n## Workflow Logic Modules\n\n",
    ]
    for module in ["Ingest and profile inputs", "Normalize to neutral schema", "Apply rules and assumptions", "Score and prioritise outputs", "Publish action and governance views"]:
        sop.append(f"- {module}\n")
    sop.append("\n## SOP Gate\n\n")
    sop.append(f"- Current gate status: **{readiness['workflow_gate']}**\n")
    sop.append(f"- Overall readiness: **{readiness['overall_pct']}%**\n")
    if readiness["blockers"]:
        sop.append("- Blockers:\n")
        for blocker in readiness["blockers"]:
            sop.append(f"  - {blocker}\n")
    else:
        sop.append("- No blockers detected. Workflow generation can be handed to the Large-only workflow builder path.\n")

    gap_log = [
        common_header,
        "## Gap Log\n\n",
        "| Section | Required | Capture | Transcript Evidence | Approval | Next Step |\n",
        "|---|---|---|---|---|---|\n",
    ]
    if gap_rows:
        for row in gap_rows:
            gap_log.append(f"| {row['section']} | {row['required']} | {row['capture']} | {row['evidence']} | {row['approval']} | {row['next_step']} |\n")
    else:
        gap_log.append("| None | yes | Complete | Supported | Approved | Ready for workflow handoff. |\n")

    assessment = [
        common_header,
        "## Architecture Assessment\n\n",
        "The cockpit treats `03_accelerator_sop.md` as the mandatory handoff gate before workflow generation.\n\n",
        "## Readiness Scores\n\n",
        f"- Capture readiness: {readiness['capture_pct']}%\n",
        f"- Human approval readiness: {readiness['approval_pct']}%\n",
        f"- Artifact readiness: {readiness['artifact_pct']}%\n",
        f"- Overall readiness: {readiness['overall_pct']}%\n\n",
        "## Workflow Build Position\n\n",
    ]
    if readiness["workflow_gate"] == "ready":
        assessment.append("The project can move into Jon's Large-only workflow build handoff after final human review.\n")
    else:
        assessment.append("The project is not ready for workflow generation. Resolve the SOP gate blockers and regenerate the handoff prompt.\n")

    content_by_name = {
        "01_customer_discovery_conversation.md": "".join(discovery),
        "02_guided_sop_capture.md": "".join(guided),
        "03_accelerator_sop.md": "".join(sop),
        "sop_gap_log.md": "".join(gap_log),
        "sop_architecture_assessment.md": "".join(assessment),
    }
    content_by_name.update(_asset_content_by_name(manifest, common_header))

    for name, content in content_by_name.items():
        path = _artifact_path(project_dir, name)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        manifest.setdefault("artifacts", {}).setdefault(name, {})
        manifest["artifacts"][name].update(
            {
                "path": str(path),
                "relative_path": str(path.relative_to(project_dir)).replace("\\", "/"),
                "exists": True,
                "status": "generated",
                "generated_at": utc_now(),
            }
        )

    refreshed_readiness = calculate_readiness(_refresh_artifact_records(manifest), sections)
    manifest = _write_peer_review_status(manifest, refreshed_readiness)
    refreshed_readiness = calculate_readiness(_refresh_artifact_records(manifest), sections)
    manifest = generate_workflow_build_handoff(manifest, sections, refreshed_readiness)
    status = _pipeline_status(manifest, refreshed_readiness)
    status_dir.mkdir(parents=True, exist_ok=True)
    _artifact_path(project_dir, PIPELINE_STATUS_ARTIFACT).write_text(json.dumps(status, indent=2), encoding="utf-8")
    _artifact_path(project_dir, PIPELINE_LOG_ARTIFACT).write_text(
        f"# Pipeline Log\n\n- {utc_now()}: Generated cockpit document chain. Workflow gate: {refreshed_readiness['workflow_gate']}.\n",
        encoding="utf-8",
    )

    return _refresh_artifact_records(manifest)


def _copy_text_to_clipboard(text: str) -> bool:
    if os.name != "nt":
        return False
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", "Set-Clipboard -Value ([Console]::In.ReadToEnd())"],
            input=text,
            text=True,
            check=True,
            capture_output=True,
            timeout=6,
        )
        return True
    except (subprocess.SubprocessError, OSError):
        return False


def launch_codex_for_run(manifest: dict[str, Any], sections: list[Section]) -> dict[str, Any]:
    readiness = calculate_readiness(_refresh_artifact_records(manifest), sections)
    if readiness["workflow_gate"] != "ready":
        raise RuntimeError("Workflow build is still blocked by the SOP gate.")

    manifest = generate_workflow_build_handoff(manifest, sections, readiness)
    workflow_build = manifest.get("workflow_build", {})
    codex_path = workflow_build.get("codex_path")
    prompt_path = Path(workflow_build.get("prompt_path", ""))
    if not codex_path:
        raise RuntimeError("Codex was not detected on this machine.")
    if not prompt_path.exists():
        raise RuntimeError("Workflow build prompt is missing. Regenerate the handoff first.")

    copied = _copy_text_to_clipboard(prompt_path.read_text(encoding="utf-8"))
    creationflags = subprocess.CREATE_NEW_CONSOLE if os.name == "nt" else 0
    subprocess.Popen([codex_path], cwd=str(run_dir(manifest["run_id"])), creationflags=creationflags)
    manifest.setdefault("workflow_build", {}).update(
        {
            "last_launch_at": utc_now(),
            "last_launch_status": "launched_codex_prompt_copied" if copied else "launched_codex_copy_failed",
        }
    )
    return _refresh_artifact_records(manifest)


def open_project_subfolder(manifest: dict[str, Any], subfolder: str) -> None:
    allowed = {
        "project": ".",
        "start_here": "00_start_here",
        "discovery": "01_discovery",
        "sop_authoring": "02_sop_authoring",
        "workflow_build": "03_workflow_build",
        "reference_examples": "04_reference_examples",
        "sequencer": "sequencer",
        "docs": "02_sop_authoring",
        "workflows": WORKFLOW_OUTPUT_DIR,
        "status": WORKFLOW_STATUS_DIR,
        "validation": WORKFLOW_VALIDATION_DIR,
    }
    relative = allowed.get(subfolder)
    if relative is None:
        raise ValueError("Unsupported folder target.")
    target = (run_dir(manifest["run_id"]) / relative).resolve()
    project_dir = run_dir(manifest["run_id"]).resolve()
    if target != project_dir and project_dir not in target.parents:
        raise ValueError("Folder target is outside this project.")
    target.mkdir(parents=True, exist_ok=True)
    if os.name == "nt":
        subprocess.Popen(["explorer", str(target)])


def sync_generated_docs(manifest: dict[str, Any]) -> dict[str, Any]:
    STARTER_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    project_dir = run_dir(manifest["run_id"])
    for name in DOC_ARTIFACTS:
        source = _artifact_path(project_dir, name)
        if source.exists():
            shutil.copy2(source, STARTER_DOCS_DIR / name)
    manifest["sync"]["synced_at"] = utc_now()
    return manifest


def delete_run(run_id: str) -> None:
    target = run_dir(run_id).resolve()
    root = RUNS_DIR.resolve()
    if target == root or root not in target.parents:
        raise ValueError("Refusing to delete a folder outside the cockpit runs directory.")
    if target.exists():
        shutil.rmtree(target)


def artifact_cards(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    records = _refresh_artifact_records(manifest).get("artifacts", {})
    return [records[name] for name in VISIBLE_ARTIFACTS if name in records]


def artifact_file_path(run_id: str, artifact_key: str) -> Path:
    manifest = load_manifest(run_id)
    records = _refresh_artifact_records(manifest).get("artifacts", {})
    key = artifact_key.replace("\\", "/")
    record = records.get(key)
    if not record or not record.get("exists"):
        raise FileNotFoundError(f"Artifact not found: {artifact_key}")
    path = Path(record["path"]).resolve()
    project_dir = run_dir(run_id).resolve()
    if project_dir != path and project_dir not in path.parents:
        raise ValueError("Artifact path is outside the project folder.")
    return path


def process_stages() -> list[dict[str, Any]]:
    return load_sequence_config().get("stages", [])


def sections_as_dicts(sections: list[Section]) -> list[dict[str, Any]]:
    return [asdict(section) for section in sections]
