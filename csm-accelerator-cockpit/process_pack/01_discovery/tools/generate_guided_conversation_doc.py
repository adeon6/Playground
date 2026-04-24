from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = r"process_pack\\guided_discovery_conversation_template.docx"


SECTIONS = [
    {
        "title": "0. Opening And Framing",
        "prompts": [
            "Today I want to understand the business problem, the current process, the desired outcome, the information needed to support it, the rules that matter, and what the first useful slice should prove.",
            "As we go, we will separate confirmed facts from assumptions and open questions.",
        ],
        "capture": [
            "Customer / engagement",
            "Primary business owner",
            "Supporting SMEs",
            "Intended outcome of discovery",
            "Whether the goal is proof of concept, first operational slice, or fuller delivery design",
        ],
    },
    {
        "title": "1. Business Problem",
        "prompts": [
            "What problem are you trying to solve?",
            "What is slow, manual, inconsistent, risky, or hard to trust today?",
            "What happens when this problem is not solved well?",
            "Who feels the pain most directly?",
            "Are you trying to improve visibility, speed, quality, prioritisation, decision-making, or something else?",
        ],
        "capture": [
            "Core business problem",
            "Current pain",
            "Consequences / business impact",
            "Affected users or teams",
            "Desired improvement theme",
        ],
    },
    {
        "title": "2. Current Process",
        "prompts": [
            "How is this handled today?",
            "What are the main steps from input to outcome?",
            "Which systems, files, reports, or teams are involved?",
            "Where do manual work, spreadsheets, emails, or handoffs happen?",
            "Which steps are most painful or fragile?",
            "What makes the current output difficult to trust or use?",
        ],
        "capture": [
            "Current-state process",
            "Systems and teams involved",
            "Manual steps and handoffs",
            "Bottlenecks",
            "Trust / quality issues",
        ],
    },
    {
        "title": "3. Desired Outcome",
        "prompts": [
            "If we gave you a useful first version quickly, what would it produce?",
            "What outputs would be most valuable first?",
            "Who would use those outputs?",
            "What action or decision should those outputs support?",
            "What does success look like for the first slice?",
            "What can wait until later?",
        ],
        "capture": [
            "Desired outputs",
            "Intended users / consumers",
            "Decisions or actions supported",
            "First-slice success criteria",
            "Deferred items",
        ],
    },
    {
        "title": "4. Business Questions",
        "prompts": [
            "What questions do you need this solution to answer?",
            "What should someone be able to decide once they see the results?",
            "Which cases should be prioritised?",
            "Which cases should be flagged for review?",
            "What patterns, exceptions, or risks matter most?",
        ],
        "capture": [
            "Core business questions",
            "Decision points",
            "Prioritisation needs",
            "Review / exception needs",
            "Business significance of answers",
        ],
    },
    {
        "title": "5. Scope",
        "prompts": [
            "What is in scope for the first slice?",
            "What is out of scope for now?",
            "Are there specific business units, regions, products, customers, channels, or process steps to include or exclude?",
            "Is there a preferred pilot boundary?",
            "Are we proving feasibility, value, or both?",
        ],
        "capture": [
            "In-scope entities / domains",
            "Out-of-scope areas",
            "Pilot boundary",
            "Phase 1 intent",
            "Explicit exclusions",
        ],
    },
    {
        "title": "6. Inputs, Sources, And Ownership",
        "prompts": [
            "What information is needed to answer these questions?",
            "Where does that information live today?",
            "Which systems, files, owners, or teams provide it?",
            "What do we know about the source structure already?",
            "What still needs to be confirmed before build?",
            "Are there known quality, access, or timing concerns?",
        ],
        "capture": [
            "Required inputs",
            "Source systems / files",
            "Data owners / SMEs",
            "Known structure and availability",
            "Unknowns requiring discovery",
            "Quality / access risks",
        ],
    },
    {
        "title": "7. Rules, Logic, And Definitions",
        "prompts": [
            "What business rules are already trusted?",
            "Which metrics, classifications, or thresholds matter?",
            "How do you currently define a good, bad, at-risk, or priority case?",
            "Are there rules for inclusion, exclusion, or segmentation?",
            "Which edge cases need special treatment?",
            "What assumptions are acceptable, and which require confirmation?",
        ],
        "capture": [
            "Trusted rules",
            "Important definitions",
            "Thresholds / classifications",
            "Inclusion / exclusion rules",
            "Edge cases",
            "Areas where logic is still uncertain",
        ],
    },
    {
        "title": "8. Exceptions And Safe Handling",
        "prompts": [
            "What should happen when data is missing, inconsistent, late, or invalid?",
            "What failure cases do you already expect?",
            "Should questionable records be excluded, flagged, defaulted, or sent for review?",
            "What outcomes would create risk if handled incorrectly?",
        ],
        "capture": [
            "Expected failure modes",
            "Safe handling rules",
            "Review / escalation paths",
            "Risk-sensitive cases",
        ],
    },
    {
        "title": "9. Validation And Trust",
        "prompts": [
            "What would make you trust the first slice?",
            "How would you validate whether it is working?",
            "What should it reconcile against?",
            "What should be explainable to the business?",
            "What would count as failure?",
            "Do different audiences need different outputs from the same logic?",
        ],
        "capture": [
            "Validation criteria",
            "Reconciliation targets",
            "Trust signals",
            "Explainability needs",
            "Failure conditions",
            "Audience-specific output needs",
        ],
    },
    {
        "title": "10. Operational Readiness And Phasing",
        "prompts": [
            "What needs to be true for this to be useful in the short term?",
            "What can be manual in phase 1?",
            "What needs to be productionised later?",
            "Are there scheduling, alerting, write-back, governance, or access requirements that should be deferred?",
            "What is the smallest slice that would still prove the approach?",
        ],
        "capture": [
            "Immediate needs",
            "Deferred operational requirements",
            "Phase 1 vs later-state expectations",
            "Minimum viable first slice",
        ],
    },
    {
        "title": "11. Close And Playback",
        "prompts": [
            "Let me summarise what I think the first slice is.",
            "Does this reflect the problem, scope, and expected outputs correctly?",
            "What have we confirmed?",
            "What are we still assuming?",
            "What must be answered before build starts?",
            "What would make this worth continuing?",
        ],
        "capture": [
            "Agreed first-slice objective",
            "Confirmed scope and outputs",
            "Outstanding questions",
            "Discovery actions",
            "Approval to proceed or refine",
        ],
    },
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bullets(cell, items):
    cell.text = ""
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.size = Pt(10)


def add_capture_block(cell, items):
    cell.text = ""
    base = [
        "Confirmed:",
        "Assumed:",
        "Unknown / To Discover:",
        "Why it matters:",
        "Next action:",
        "",
        "Specific capture for this section:",
    ]
    for line in base:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.size = Pt(10)
        if line.endswith(":") or line == "Specific capture for this section:":
            run.bold = True
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.size = Pt(10)
    cell.add_paragraph("")


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(20)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Guided Discovery Conversation Template")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run(
    "Industry- and problem-agnostic template for consultant-led discovery, SOP capture, and first-slice workflow scoping."
)
subtitle_run.font.size = Pt(10)

intro = doc.add_paragraph()
intro_run = intro.add_run(
    "Use the left column for live prompts and the right column for structured capture. "
    "For every section, separate confirmed facts from assumptions and open questions."
)
intro_run.font.size = Pt(10)

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
table.autofit = False

header_cells = table.rows[0].cells
header_cells[0].width = Inches(4.7)
header_cells[1].width = Inches(5.8)
for idx, label in enumerate(("Consultant Prompts", "Capture")):
    cell = header_cells[idx]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    shade_cell(cell, "D9EAF7")
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

for section_data in SECTIONS:
    section_row = table.add_row().cells
    merged = section_row[0].merge(section_row[1])
    merged.text = ""
    p = merged.paragraphs[0]
    run = p.add_run(section_data["title"])
    run.bold = True
    run.font.size = Pt(11)
    shade_cell(merged, "EAF3E2")
    set_cell_margins(merged)

    row_cells = table.add_row().cells
    row_cells[0].width = Inches(4.7)
    row_cells[1].width = Inches(5.8)
    for cell in row_cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_bullets(row_cells[0], section_data["prompts"])
    add_capture_block(row_cells[1], section_data["capture"])

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
